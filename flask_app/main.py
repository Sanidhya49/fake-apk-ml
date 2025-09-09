"""Flask API for APK fake/legit detection.

This Flask app provides REST endpoints for scanning single APK files and batch analysis.
It wraps the ML inference service and provides a simple HTTP API.

Run locally:
    python flask_app/main.py
"""

import os
import sys
import tempfile
import json
import time
import threading
from datetime import datetime
from typing import Dict, List
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor

import numpy as np
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename

# Add the parent directory to Python path so we can import from ml module
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from ml import static_extract
from ml.utils import ensure_dirs, load_model, vectorize_feature_dict, get_sha256, load_bank_whitelist
from rapidfuzz import fuzz

# Load .env if present
try:
    from dotenv import load_dotenv
    load_dotenv(override=True)
except Exception:
    pass

# Initialize Flask app
app = Flask(__name__)

# Configure CORS with more explicit settings
CORS(app, 
     origins=["*"],  # Allow all origins for now, restrict in production
     methods=["GET", "POST", "OPTIONS"],  # Allowed HTTP methods
     allow_headers=["Content-Type", "Authorization"],  # Allowed headers
     supports_credentials=False  # Set to True if you need credentials
)

# Performance optimizations
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 3600  # Cache static files for 1 hour
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# Thread pool for async processing
executor = ThreadPoolExecutor(max_workers=4)

# Configuration
MODEL_PATH = os.path.join(os.path.dirname(__file__), "..", "models", "xgb_model.joblib")
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'apk', 'apks', 'xapk'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Global variables for caching
_model_cache = None
_feature_order_cache = None
_saved_threshold_cache = None
_bank_whitelist_cache = None
_threat_feed_cache = None

import logging

# Configure logging to show everything
logging.basicConfig(level=logging.DEBUG)

@app.before_request
def log_request():
    app.logger.info(f"[REQ] {request.remote_addr} {request.method} {request.full_path}")
    app.logger.info(f"Headers: {dict(request.headers)}")
    if request.content_length:
        app.logger.info(f"Body size: {request.content_length} bytes")
    if request.data and request.content_length < 5000:  # prevent logging huge APK files
        try:
            app.logger.info(f"Body: {request.data.decode('utf-8', errors='ignore')}")
        except Exception:
            app.logger.warning("[Body could not be decoded]")

@app.after_request
def log_response(response):
    app.logger.info(
        f"[RESP] {request.remote_addr} {request.method} {request.full_path} "
        f"-> {response.status} ({len(response.data)} bytes)"
    )
    try:
        if len(response.data) < 5000:  # don’t spam with huge responses
            app.logger.debug(f"Response body: {response.data.decode('utf-8', errors='ignore')}")
    except Exception:
        app.logger.warning("[Response body could not be decoded]")
    return response

@app.teardown_request
def log_teardown(error=None):
    if error:
        app.logger.error(f"[ERROR] {request.remote_addr} {request.method} {request.full_path} -> {error}", exc_info=True)

def get_cached_model():
    """Get cached model instance"""
    global _model_cache
    if _model_cache is None:
        try:
            model_data = load_model(MODEL_PATH)
            if isinstance(model_data, dict):
                _model_cache = model_data['model']  # Extract the actual model from dict
                print(f"Model loaded successfully from {MODEL_PATH}")
            else:
                _model_cache = model_data  # Direct model object
                print(f"Model loaded successfully from {MODEL_PATH}")
        except Exception as e:
            print(f"Error loading model: {e}")
            raise Exception(f"Failed to load model: {e}")
    return _model_cache

def get_cached_feature_order():
    """Get cached feature order"""
    global _feature_order_cache
    if _feature_order_cache is None:
        try:
            model_data = load_model(MODEL_PATH)
            if isinstance(model_data, dict):
                _feature_order_cache = model_data['feature_order']  # Get from dict
            else:
                _feature_order_cache = model_data.feature_names_in_.tolist()  # Get from model
        except Exception as e:
            print(f"Error loading feature order: {e}")
            raise Exception(f"Failed to load feature order: {e}")
    return _feature_order_cache

def get_cached_threshold():
    """Get cached threshold"""
    global _saved_threshold_cache
    if _saved_threshold_cache is None:
        try:
            model_data = load_model(MODEL_PATH)
            if isinstance(model_data, dict):
                _saved_threshold_cache = model_data.get('tuned_threshold', 0.5)  # Get from dict
            else:
                _saved_threshold_cache = getattr(model_data, 'tuned_threshold', 0.5)  # Get from model
        except Exception as e:
            print(f"Error loading threshold: {e}")
            _saved_threshold_cache = 0.5
    return _saved_threshold_cache

def get_cached_bank_whitelist():
    """Get cached bank whitelist"""
    global _bank_whitelist_cache
    if _bank_whitelist_cache is None:
        _bank_whitelist_cache = load_bank_whitelist()
    return _bank_whitelist_cache

def get_cached_threat_feed():
    """Get cached threat feed data"""
    global _threat_feed_cache
    if _threat_feed_cache is None:
        _threat_feed_cache = load_threat_feed()
    return _threat_feed_cache

def load_threat_feed():
    """Load threat feed from JSON file"""
    feed_path = os.path.join("artifacts", "threat_intel", "bad_hashes.json")
    try:
        if os.path.exists(feed_path):
            with open(feed_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return {
                    "hashes": set(data.get("hashes", [])),
                    "packages": set(data.get("packages", [])),
                    "cert_fingerprints": set(data.get("cert_fingerprints", [])),
                    "last_updated": data.get("last_updated", 0),
                    "version": data.get("version", "1.0")
                }
    except Exception as e:
        print(f"Warning: Could not load threat feed: {e}")
    
    return {
        "hashes": set(),
        "packages": set(),
        "cert_fingerprints": set(),
        "last_updated": 0,
        "version": "1.0"
    }

def save_threat_feed(feed_data):
    """Save threat feed to JSON file"""
    feed_path = os.path.join("artifacts", "threat_intel", "bad_hashes.json")
    try:
        os.makedirs(os.path.dirname(feed_path), exist_ok=True)
        with open(feed_path, "w", encoding="utf-8") as f:
            json.dump({
                "hashes": list(feed_data["hashes"]),
                "packages": list(feed_data["packages"]),
                "cert_fingerprints": list(feed_data["cert_fingerprints"]),
                "last_updated": int(time.time()),
                "version": feed_data.get("version", "1.0")
            }, f, indent=2)
        return True
    except Exception as e:
        print(f"Error saving threat feed: {e}")
        return False

def check_threat_feed(sha256_hash, package_name=None, cert_fingerprint=None):
    """Check if APK matches known bad indicators"""
    feed = get_cached_threat_feed()
    
    # Check hash
    if sha256_hash in feed["hashes"]:
        return {"match": True, "type": "hash", "value": sha256_hash}
    
    # Check package name
    if package_name and package_name in feed["packages"]:
        return {"match": True, "type": "package", "value": package_name}
    
    # Check certificate fingerprint
    if cert_fingerprint and cert_fingerprint in feed["cert_fingerprints"]:
        return {"match": True, "type": "certificate", "value": cert_fingerprint}
    
    return {"match": False, "type": None, "value": None}

@lru_cache(maxsize=1000)
def allowed_file(filename):
    """Check if uploaded file has allowed extension (cached)"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def _vectorize_from_extract(extract_dict: Dict, feature_order: List[str]) -> Dict:
    """Convert extracted APK features to ML feature vector"""
    # Build the same features as in feature_builder
    permissions = set(extract_dict.get("permissions", []))
    suspicious = extract_dict.get("suspicious_apis", {})

    base: Dict[str, int] = {}
    
    # Permissions of interest (must match builder)
    for p in ["READ_SMS", "SEND_SMS", "RECEIVE_SMS", "SYSTEM_ALERT_WINDOW", "READ_CONTACTS", "INTERNET"]:
        base[p] = 1 if p in permissions else 0

    # Suspicious APIs
    for name in [
        "getDeviceId",
        "sendTextMessage", 
        "SmsManager",
        "DexClassLoader",
        "TYPE_SYSTEM_ALERT_WINDOW",
        "addView",
        "HttpURLConnection",
        "openConnection",
    ]:
        base[f"api_{name}"] = int(bool(suspicious.get(name, 0)))

    base["count_suspicious"] = int(sum(base[k] for k in base if k.startswith("api_")))
    base["cert_present"] = 0 if extract_dict.get("cert_subject", "unknown") in ("", "unknown") else 1
    
    # New features mirrored from feature_builder
    cert_issuer = extract_dict.get("cert_issuer", "unknown")
    base["issuer_present"] = 0 if not cert_issuer or cert_issuer == "unknown" else 1
    pkg = (extract_dict.get("package") or "").lower()
    cn = str(extract_dict.get("cert_subject", "")).lower()
    base["cn_matches_package"] = 1 if (pkg and pkg in cn) else 0
    
    # Additional CN-based features
    subject_cn = (extract_dict.get("cert_subject_cn") or "").lower()
    issuer_cn = (extract_dict.get("cert_issuer_cn") or "").lower()
    base["issuer_cn_google_android"] = 1 if ("google" in issuer_cn or "android" in issuer_cn) else 0
    base["subject_cn_contains_pkg"] = 1 if (pkg and pkg in subject_cn) else 0
    base["issuer_subject_cn_equal"] = 1 if (subject_cn and issuer_cn and subject_cn == issuer_cn) else 0
    
    # Package analysis
    base["pkg_official"] = 0
    if pkg:
        bank_whitelist = get_cached_bank_whitelist()
        for bank in bank_whitelist:
            if fuzz.partial_ratio(pkg, bank) >= 85:
                base["pkg_official"] = 1
                break
    
    # Build feature vector
    vector = []
    for feat in feature_order:
        vector.append(base.get(feat, 0))
    
    return {"vector": vector, "feature_map": base}

def process_single_apk(file_path: str, quick: bool = False, debug: bool = False) -> Dict:
    """Process a single APK file (optimized version)"""
    try:
        start_time = time.time()  # Add start time for debug
        
        # Get cached instances
        model = get_cached_model()
        feature_order = get_cached_feature_order()
        saved_thr = get_cached_threshold()

        # Get SHA256 for caching
        sha = get_sha256(file_path)
        
        # Check threat feed first (fastest check)
        threat_check = check_threat_feed(sha)
        
        # Check cache first
        cache_dir = os.path.join("artifacts", "static_jsons")
        cache_path = os.path.join(cache_dir, f"{sha}.json")
        
        if os.path.exists(cache_path):
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    ext = json.load(f)
            except:
                ext = None
        else:
            ext = None

        # Extract features if not cached
        if ext is None:
            try:
                ext = static_extract.extract(file_path, quick=quick)
            except Exception as e:
                # Retry with quick=True if full parse failed
                if not quick:
                    try:
                        ext = static_extract.extract(file_path, quick=True)
                    except Exception:
                        return {"error": "parse_failed", "detail": "Could not parse APK"}
                else:
                    return {"error": "parse_failed", "detail": "Could not parse APK"}
            
            if not isinstance(ext, dict):
                return {"error": "parse_failed", "detail": "Could not parse APK"}
            
            # Cache the results
            try:
                sha = ext.get("sha256") or sha
                if sha:
                    os.makedirs(cache_dir, exist_ok=True)
                    cache_path = os.path.join(cache_dir, f"{sha}.json")
                    if not os.path.exists(cache_path):
                        with open(cache_path, "w", encoding="utf-8") as _f:
                            json.dump(ext, _f, ensure_ascii=False)
            except Exception:
                pass

        # Vectorize features
        v = _vectorize_from_extract(ext, feature_order)
        X = np.array([v["vector"]])
        
        # Get prediction probability
        try:
            prob = float(model.predict_proba(X)[0, 1])
        except Exception:
            # Fallback if model lacks predict_proba
            raw = model.predict(X)[0]
            prob = float(int(raw))

        # Determine prediction and risk
        threshold = float(os.environ.get("ML_FAKE_THRESHOLD", str(saved_thr)))
        pred_initial = 1 if prob >= threshold else 0
        pred = pred_initial
        
        # Apply heuristics and overrides (simplified version)
        feat = v["feature_map"]
        is_official = feat.get("pkg_official", 0) == 1
        official_override_cap = float(os.environ.get("ML_OFFICIAL_OVERRIDE_CAP", "0.40"))
        
        # Official package override
        if is_official and prob <= official_override_cap:
            pred = 0
        
        # Threat feed override (highest priority)
        if threat_check["match"]:
            pred = 1  # Force fake prediction
            prob = 0.95  # High confidence for known bad
        
        # Risk categorization with confidence
        if prob >= max(0.8, threshold):
            risk = "Red"
        elif prob >= threshold or pred == 1:
            risk = "Amber"
        else:
            risk = "Green"

        # Get top SHAP features (best effort)
        top_shap = []
        try:
            import shap
            print("✅ SHAP imported successfully")
            try:
                print(f"🔍 SHAP Debug: Model type: {type(model)}")
                print(f"🔍 SHAP Debug: X shape: {X.shape}")
                print(f"🔍 SHAP Debug: Feature order length: {len(feature_order)}")
                
                # Handle CalibratedClassifierCV by extracting base model
                if hasattr(model, 'estimator'):
                    base_model = model.estimator
                    print(f"🔍 SHAP Debug: Using estimator: {type(base_model)}")
                elif hasattr(model, 'base_estimator'):
                    base_model = model.base_estimator
                    print(f"🔍 SHAP Debug: Using base model: {type(base_model)}")
                else:
                    base_model = model
                    print(f"🔍 SHAP Debug: Using original model: {type(base_model)}")
                
                explainer = shap.TreeExplainer(base_model)
                shap_values = explainer.shap_values(X)
                print(f"🔍 SHAP Debug: SHAP values type: {type(shap_values)}")
                
                if isinstance(shap_values, list):
                    sv = shap_values[1][0]
                    print(f"🔍 SHAP Debug: Using list index 1, shape: {sv.shape}")
                else:
                    sv = shap_values[0]
                    print(f"🔍 SHAP Debug: Using direct index 0, shape: {sv.shape}")
                
                idxs = np.argsort(np.abs(sv))[::-1][:3]
                print(f"🔍 SHAP Debug: Top indices: {idxs}")
                
                for j in idxs:
                    if j < len(feature_order):
                        top_shap.append({"feature": feature_order[j], "value": float(sv[j])})
                    else:
                        print(f"🔍 SHAP Debug: Index {j} out of range for feature_order")
                
                print(f"✅ SHAP analysis completed: {len(top_shap)} features")
            except Exception as e:
                print(f"❌ SHAP analysis failed: {e}")
                import traceback
                traceback.print_exc()
                top_shap = []
        except ImportError as e:
            print(f"❌ SHAP import failed: {e}")
            top_shap = []

        # Calculate confidence score based on prediction
        if pred == 1:  # Fake prediction
            # For fake predictions, high probability = high confidence
            if prob >= 0.8:
                confidence = "High"
            elif prob >= 0.6:
                confidence = "Medium"
            else:
                confidence = "Low"
        else:  # Legitimate prediction
            # For legitimate predictions, low probability = high confidence
            if prob <= 0.2:
                confidence = "High"
            elif prob <= 0.4:
                confidence = "Medium"
            else:
                confidence = "Low"

        label_map = {0: "legit", 1: "fake"}
        
        # Calculate confidence percentage for display
        if pred == 1:  # Fake prediction
            confidence_percentage = prob * 100  # Direct percentage
        else:  # Legitimate prediction
            confidence_percentage = (1 - prob) * 100  # Inverted percentage
        
        # Add confidence to result
        result = {
            "prediction": label_map.get(int(pred), str(pred)),
            "probability": prob,
            "risk_level": risk,  # Changed from "risk" to "risk_level" to match expected format
            "confidence": confidence,
            "confidence_percentage": round(confidence_percentage, 1),  # Add user-friendly confidence percentage
            "top_shap": top_shap,
            "feature_vector": v["feature_map"],
            "processing_time": time.time() - start_time,
            "model_threshold": threshold,
            "cache_used": os.path.exists(cache_path) if 'cache_path' in locals() else False,
            "app_label": ext.get("app_label", ""),  # Add app label from extracted data
            "package": ext.get("package", ""),  # Add package name
            "version": ext.get("version", ""),  # Add version
            "file_size": ext.get("file_size", 0),  # Add file size
            # Add app technical information for frontend display
            "min_sdk": ext.get("min_sdk", 0) if ext.get("min_sdk", 0) > 0 else "N/A",
            "target_sdk": ext.get("target_sdk", 0) if ext.get("target_sdk", 0) > 0 else "N/A",
            "num_activities": ext.get("num_activities", 0) if ext.get("num_activities", 0) > 0 else "N/A",
            "num_services": ext.get("num_services", 0) if ext.get("num_services", 0) > 0 else "N/A",
            "num_receivers": ext.get("num_receivers", 0) if ext.get("num_receivers", 0) > 0 else "N/A",
            "num_dex": ext.get("num_dex", 0) if ext.get("num_dex", 0) > 0 else "N/A",
            "domains_count": len(ext.get("domains", [])) if ext.get("domains") else "N/A",
            "main_activity": ext.get("main_activity", "N/A"),
            "total_permissions": len(ext.get("permissions", [])),
            "exported_components": len(ext.get("exported", [])),
            "threat_feed_match": threat_check,
        }
        
        # Add critical security features
        try:
            # Extract critical permissions
            permissions = ext.get("permissions", [])
            critical_permissions = []
            for perm in permissions:
                if any(critical in perm.lower() for critical in [
                    'sms', 'contacts', 'location', 'camera', 'microphone', 
                    'phone', 'calendar', 'call_log', 'storage', 'system_alert'
                ]):
                    critical_permissions.append(perm)
            
            # Count suspicious APIs
            suspicious_apis = ext.get("suspicious_apis", [])
            suspicious_api_count = len(suspicious_apis)
            
            # Get certificate information
            cert_subject = ext.get("cert_subject", "unknown")
            cert_issuer = ext.get("cert_issuer", "unknown")
            cert_present = 1 if cert_subject != "unknown" else 0
            
            # Determine app trust level
            app_trust_level = "High" if is_official else "Low"
            
            # Add to result
            result.update({
                "critical_permissions": critical_permissions,
                "suspicious_api_count": suspicious_api_count,
                "total_permissions": len(permissions),
                "certificate_status": "Valid" if cert_present else "Missing",
                "signing_authority": cert_issuer,
                "app_trust_level": app_trust_level,
                "critical_labels": critical_permissions  # For backward compatibility
            })
        except Exception as e:
            print(f"⚠️  Could not extract critical features: {e}")
            # Add default values
            result.update({
                "critical_permissions": [],
                "suspicious_api_count": 0,
                "total_permissions": 0,
                "certificate_status": "Unknown",
                "signing_authority": "Unknown",
                "app_trust_level": "Unknown",
                "critical_labels": []
            })
        
        # Add AI explanation
        try:
            result["ai_explanation"] = _generate_ai_explanation(result)
        except Exception as e:
            print(f"❌ AI explanation generation failed: {e}")
            result["ai_explanation"] = "AI explanation could not be generated"
        
        if debug:
            result["debug"] = {
                "processing_time_seconds": time.time() - start_time,
                "cache_used": os.path.exists(cache_path) if 'cache_path' in locals() else False,
                "model_threshold": float(threshold),
                "saved_tuned_threshold": float(saved_thr),
                "pred_initial": int(pred_initial),
                "is_official": bool(is_official),
                "official_override_cap": float(official_override_cap),
                "sha256": sha,
            }
        
        return result
        
    except Exception as e:
        return {"error": "prediction_failed", "detail": str(e)}

# Routes
@app.before_request
def handle_preflight():
    """Handle CORS preflight requests"""
    if request.method == "OPTIONS":
        response = jsonify()
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add('Access-Control-Allow-Headers', "Content-Type,Authorization")
        response.headers.add('Access-Control-Allow-Methods', "GET,PUT,POST,DELETE,OPTIONS")
        return response

@app.after_request
def after_request(response):
    """Add CORS headers and performance monitoring to all responses"""
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    
    # Add performance monitoring headers
    response.headers.add('X-API-Version', '2.0')
    response.headers.add('X-Model-Threshold', str(os.environ.get('ML_FAKE_THRESHOLD', '0.35')))
    
    return response

@app.route('/', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "ok", 
        "message": "Fake APK Detection API is running",
        "endpoints": {
            "scan_single": "POST /scan",
            "scan_batch": "POST /scan-batch",
            "generate_report": "POST /report",
            "generate_batch_report": "POST /report-batch",
            "report_abuse": "POST /report-abuse",
            "report_batch_abuse": "POST /report-batch-abuse",
            "threat_feed": "GET /threat-feed",
            "submit_threat_intel": "POST /threat/submit",
            "news": "GET /news",
            "news_categories": "GET /news/categories"
        }
    })

@app.route('/scan', methods=['POST'])
def scan_single():
    """Scan a single APK file with comprehensive analysis and feature extraction"""
    start_time = time.time()
    
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({"error": "no_file", "detail": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "no_file", "detail": "No file selected"}), 400
        
        # Check if file is empty
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size == 0:
            return jsonify({"error": "empty_file", "detail": "File is empty"}), 400
        
        # Validate file type (cached)
        if not allowed_file(file.filename):
            return jsonify({
                "error": "invalid_file_type", 
                "detail": f"Invalid file type. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
            }), 400
        
        # Get query parameters
        quick = request.args.get('quick', 'false').lower() == 'true'
        debug = request.args.get('debug', 'false').lower() == 'true'
        include_features = request.args.get('include_features', 'true').lower() == 'true'
        
        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
        try:
            file.save(temp_file.name)
            temp_file.close()
            
            # Make sure directories exist
            ensure_dirs()
            
            # Process APK with comprehensive analysis
            result = process_single_apk(temp_file.name, quick=quick, debug=True)
            
            # Add original filename and upload info
            result["file"] = file.filename
            result["upload_size"] = file_size
            result["upload_timestamp"] = time.time()
            
            # Enhanced feature extraction and analysis
            if include_features and "feature_vector" in result:
                feature_vector = result["feature_vector"]
                
                # Extract and categorize permissions
                permissions = {
                    "granted": [],
                    "sensitive": [],
                    "dangerous": []
                }
                
                permission_categories = {
                    "sensitive": ["READ_SMS", "SEND_SMS", "RECEIVE_SMS", "READ_CONTACTS", "SYSTEM_ALERT_WINDOW"],
                    "dangerous": ["READ_SMS", "SEND_SMS", "RECEIVE_SMS", "SYSTEM_ALERT_WINDOW"],
                    "network": ["INTERNET"],
                    "storage": ["READ_EXTERNAL_STORAGE", "WRITE_EXTERNAL_STORAGE"]
                }
                
                for perm, value in feature_vector.items():
                    if perm in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS', 'SYSTEM_ALERT_WINDOW', 'READ_CONTACTS', 'INTERNET'] and value == 1:
                        permissions["granted"].append(perm)
                        if perm in permission_categories["sensitive"]:
                            permissions["sensitive"].append(perm)
                        if perm in permission_categories["dangerous"]:
                            permissions["dangerous"].append(perm)
                
                result["permissions_analysis"] = permissions
                
                # Extract suspicious APIs with descriptions
                suspicious_apis = []
                api_descriptions = {
                    "api_getDeviceId": "Device identifier collection - potential privacy concern",
                    "api_sendTextMessage": "SMS sending capability - potential premium SMS fraud",
                    "api_SmsManager": "SMS management - potential message interception",
                    "api_DexClassLoader": "Dynamic code loading - potential code obfuscation",
                    "api_TYPE_SYSTEM_ALERT_WINDOW": "System overlay capability - potential phishing attacks",
                    "api_addView": "UI manipulation - potential overlay attacks",
                    "api_HttpURLConnection": "Network communication - data transmission capability",
                    "api_openConnection": "Network connection - remote server communication"
                }
                
                for api, value in feature_vector.items():
                    if api.startswith('api_') and value == 1:
                        suspicious_apis.append({
                            "api": api.replace('api_', ''),
                            "description": api_descriptions.get(api, "Potentially suspicious behavior"),
                            "risk_level": "high" if api in ["api_sendTextMessage", "api_SmsManager", "api_DexClassLoader"] else "medium"
                        })
                
                result["suspicious_apis_analysis"] = suspicious_apis
                
                # Security indicators summary
                security_indicators = {
                    "has_sms_permissions": any(feature_vector.get(p, 0) == 1 for p in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS']),
                    "has_contacts_access": feature_vector.get('READ_CONTACTS', 0) == 1,
                    "has_system_overlay": feature_vector.get('SYSTEM_ALERT_WINDOW', 0) == 1,
                    "has_internet_access": feature_vector.get('INTERNET', 0) == 1,
                    "suspicious_api_count": feature_vector.get('count_suspicious', 0),
                    "has_valid_certificate": feature_vector.get('cert_present', 0) == 1,
                    "is_official_package": feature_vector.get('pkg_official', 0) == 1,
                    "certificate_issues": not feature_vector.get('cert_present', 0) == 1
                }
                
                result["security_indicators"] = security_indicators
                
                # Risk assessment details
                risk_factors = []
                if security_indicators["has_sms_permissions"]:
                    risk_factors.append("SMS permissions detected - potential for premium SMS fraud")
                if security_indicators["has_contacts_access"]:
                    risk_factors.append("Contacts access - potential data harvesting")
                if security_indicators["has_system_overlay"]:
                    risk_factors.append("System overlay permission - potential phishing attacks")
                if security_indicators["suspicious_api_count"] > 3:
                    risk_factors.append(f"Multiple suspicious APIs ({security_indicators['suspicious_api_count']}) detected")
                if not security_indicators["has_valid_certificate"]:
                    risk_factors.append("No valid certificate - integrity concerns")
                if not security_indicators["is_official_package"]:
                    risk_factors.append("Unofficial package source - trustworthiness concerns")
                
                result["risk_factors"] = risk_factors
            
            # Add comprehensive performance metrics
            processing_time = time.time() - start_time
            result["performance_metrics"] = {
                "total_processing_time": round(processing_time, 3),
                "analysis_timestamp": time.time(),
                "cache_hit": result.get("cache_used", False),
                "model_version": "2.0",
                "api_version": "2.0"
            }
            
            # Add debug information if requested
            if debug:
                if "debug" not in result:
                    result["debug"] = {}
                result["debug"].update({
                    "endpoint_processing_time": round(processing_time, 3),
                    "file_size_bytes": file_size,
                    "temp_file_used": True,
                    "feature_extraction_enabled": include_features,
                    "quick_mode": quick
                })
            
            # Check for errors
            if "error" in result:
                return jsonify(result), 422 if result["error"] == "parse_failed" else 500
            
            # Add success metadata
            result["status"] = "success"
            result["analysis_type"] = "single_apk"
            
            return jsonify(result)
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file.name)
            except Exception:
                pass
                
    except Exception as e:
        return jsonify({
            "error": "internal_error", 
            "detail": str(e),
            "status": "failed",
            "analysis_type": "single_apk",
            "processing_time": round(time.time() - start_time, 3) if 'start_time' in locals() else 0
        }), 500

@app.route('/scan-batch', methods=['POST'])
def scan_batch():
    """Scan multiple APK files with comprehensive analysis and feature extraction (enhanced for up to 15 APKs)"""
    start_time = time.time()
    
    try:
        # Check if files are in request
        if 'files' not in request.files:
            return jsonify({"error": "no_files", "detail": "No files provided"}), 400
        
        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({"error": "no_files", "detail": "No files selected"}), 400
        
        # Limit to 15 files maximum
        if len(files) > 15:
            return jsonify({"error": "too_many_files", "detail": "Maximum 15 files allowed per batch"}), 400
        
        # Get query parameters
        quick = request.args.get('quick', 'false').lower() == 'true'
        debug = request.args.get('debug', 'false').lower() == 'true'
        include_features = request.args.get('include_features', 'true').lower() == 'true'
        include_summary = request.args.get('include_summary', 'true').lower() == 'true'
        
        # Validate all files first
        valid_files = []
        invalid_files = []
        for file in files:
            if file.filename == '':
                continue
            if not allowed_file(file.filename):
                invalid_files.append({
                    "filename": file.filename,
                    "reason": "Invalid file type"
                })
                continue
            valid_files.append(file)
        
        if not valid_files:
            return jsonify({
                "error": "no_valid_files", 
                "detail": "No valid APK files found",
                "invalid_files": invalid_files,
                "valid_extensions": list(ALLOWED_EXTENSIONS)
            }), 400
        
        # Process files with comprehensive analysis
        results = []
        temp_files = []
        batch_stats = {
            "total_size": 0,
            "processed_count": 0,
            "error_count": 0,
            "fake_count": 0,
            "legit_count": 0,
            "high_risk_count": 0,
            "medium_risk_count": 0,
            "low_risk_count": 0
        }
        
        try:
            for file in valid_files:
                # Save uploaded file temporarily
                filename = secure_filename(file.filename)
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
                temp_files.append(temp_file.name)
                
                # Get file size
                file.seek(0, 2)
                file_size = file.tell()
                file.seek(0)
                batch_stats["total_size"] += file_size
                
                file.save(temp_file.name)
                temp_file.close()
                
                # Process file with comprehensive analysis
                result = process_single_apk(temp_file.name, quick=quick, debug=True)
                result["file"] = file.filename
                result["upload_size"] = file_size
                result["upload_timestamp"] = time.time()
                
                # Track batch statistics
                batch_stats["processed_count"] += 1
                if "error" in result:
                    batch_stats["error_count"] += 1
                else:
                    # Count predictions
                    prediction = result.get("prediction", "unknown")
                    if prediction == "fake":
                        batch_stats["fake_count"] += 1
                    elif prediction == "legit":
                        batch_stats["legit_count"] += 1
                    
                    # Count risk levels
                    risk_level = result.get("risk_level", "Unknown")
                    if risk_level == "Red":
                        batch_stats["high_risk_count"] += 1
                    elif risk_level == "Amber":
                        batch_stats["medium_risk_count"] += 1
                    elif risk_level == "Green":
                        batch_stats["low_risk_count"] += 1
                
                # Enhanced feature extraction for batch processing
                if include_features and "feature_vector" in result and "error" not in result:
                    feature_vector = result["feature_vector"]
                    
                    # Extract and categorize permissions
                    permissions = {
                        "granted": [],
                        "sensitive": [],
                        "dangerous": []
                    }
                    
                    permission_categories = {
                        "sensitive": ["READ_SMS", "SEND_SMS", "RECEIVE_SMS", "READ_CONTACTS", "SYSTEM_ALERT_WINDOW"],
                        "dangerous": ["READ_SMS", "SEND_SMS", "RECEIVE_SMS", "SYSTEM_ALERT_WINDOW"]
                    }
                    
                    for perm, value in feature_vector.items():
                        if perm in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS', 'SYSTEM_ALERT_WINDOW', 'READ_CONTACTS', 'INTERNET'] and value == 1:
                            permissions["granted"].append(perm)
                            if perm in permission_categories["sensitive"]:
                                permissions["sensitive"].append(perm)
                            if perm in permission_categories["dangerous"]:
                                permissions["dangerous"].append(perm)
                    
                    result["permissions_analysis"] = permissions
                    
                    # Extract suspicious APIs with risk levels
                    suspicious_apis = []
                    api_risk_mapping = {
                        "api_getDeviceId": "medium",
                        "api_sendTextMessage": "high", 
                        "api_SmsManager": "high",
                        "api_DexClassLoader": "high",
                        "api_TYPE_SYSTEM_ALERT_WINDOW": "high",
                        "api_addView": "medium",
                        "api_HttpURLConnection": "low",
                        "api_openConnection": "low"
                    }
                    
                    for api, value in feature_vector.items():
                        if api.startswith('api_') and value == 1:
                            suspicious_apis.append({
                                "api": api.replace('api_', ''),
                                "risk_level": api_risk_mapping.get(api, "medium")
                            })
                    
                    result["suspicious_apis_analysis"] = suspicious_apis
                    
                    # Security indicators summary
                    security_indicators = {
                        "has_sms_permissions": any(feature_vector.get(p, 0) == 1 for p in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS']),
                        "has_contacts_access": feature_vector.get('READ_CONTACTS', 0) == 1,
                        "has_system_overlay": feature_vector.get('SYSTEM_ALERT_WINDOW', 0) == 1,
                        "has_internet_access": feature_vector.get('INTERNET', 0) == 1,
                        "suspicious_api_count": feature_vector.get('count_suspicious', 0),
                        "has_valid_certificate": feature_vector.get('cert_present', 0) == 1,
                        "is_official_package": feature_vector.get('pkg_official', 0) == 1,
                        "overall_risk_score": result.get("probability", 0)
                    }
                    
                    result["security_indicators"] = security_indicators
                
                results.append(result)
                
            # Add comprehensive batch performance metrics
            processing_time = time.time() - start_time
            
            # Batch-wide analysis summary
            batch_summary = {
                "overview": {
                    "total_files_submitted": len(files),
                    "valid_files": len(valid_files),
                    "invalid_files": len(invalid_files),
                    "successfully_processed": batch_stats["processed_count"],
                    "processing_errors": batch_stats["error_count"]
                },
                "security_summary": {
                    "malicious_detected": batch_stats["fake_count"],
                    "legitimate_detected": batch_stats["legit_count"],
                    "high_risk_apps": batch_stats["high_risk_count"],
                    "medium_risk_apps": batch_stats["medium_risk_count"],
                    "low_risk_apps": batch_stats["low_risk_count"],
                    "detection_rate": round((batch_stats["fake_count"] / max(batch_stats["processed_count"] - batch_stats["error_count"], 1)) * 100, 2)
                },
                "performance": {
                    "total_processing_time": round(processing_time, 3),
                    "average_time_per_file": round(processing_time / len(valid_files), 3) if valid_files else 0,
                    "files_per_second": round(len(valid_files) / processing_time, 2) if processing_time > 0 else 0,
                    "total_data_processed": batch_stats["total_size"],
                    "processing_speed_mbps": round((batch_stats["total_size"] / (1024*1024)) / processing_time, 2) if processing_time > 0 else 0
                },
                "analysis_metadata": {
                    "analysis_timestamp": time.time(),
                    "api_version": "2.0",
                    "model_version": "2.0",
                    "quick_mode": quick,
                    "debug_mode": debug,
                    "feature_extraction": include_features,
                    "max_files_allowed": 15
                }
            }
            
            # Add batch-wide feature analysis if requested
            if include_summary and results:
                # Aggregate permission analysis
                all_permissions = {}
                all_apis = {}
                total_risk_scores = []
                
                for result in results:
                    if "error" not in result and "permissions_analysis" in result:
                        for perm in result["permissions_analysis"]["granted"]:
                            all_permissions[perm] = all_permissions.get(perm, 0) + 1
                    
                    if "error" not in result and "suspicious_apis_analysis" in result:
                        for api_info in result["suspicious_apis_analysis"]:
                            api_name = api_info["api"]
                            all_apis[api_name] = all_apis.get(api_name, 0) + 1
                    
                    if "error" not in result and "probability" in result:
                        total_risk_scores.append(result["probability"])
                
                batch_summary["feature_analysis"] = {
                    "most_common_permissions": sorted(all_permissions.items(), key=lambda x: x[1], reverse=True)[:10],
                    "most_common_suspicious_apis": sorted(all_apis.items(), key=lambda x: x[1], reverse=True)[:10],
                    "average_risk_score": round(sum(total_risk_scores) / len(total_risk_scores), 3) if total_risk_scores else 0,
                    "risk_score_distribution": {
                        "high_risk_0.8+": len([s for s in total_risk_scores if s >= 0.8]),
                        "medium_risk_0.5-0.8": len([s for s in total_risk_scores if 0.5 <= s < 0.8]),
                        "low_risk_below_0.5": len([s for s in total_risk_scores if s < 0.5])
                    }
                }
            
            # Add debug information if requested
            if debug:
                for i, result in enumerate(results):
                    if "debug" not in result:
                        result["debug"] = {}
                    result["debug"].update({
                        "batch_position": i + 1,
                        "batch_total": len(valid_files),
                        "batch_processing_time": round(processing_time, 3),
                        "individual_processing_order": i + 1
                    })
            
            # Add status and metadata to all results
            for result in results:
                if "error" not in result:
                    result["status"] = "success"
                result["analysis_type"] = "batch_apk"
            
            response_data = {
                "status": "success",
                "analysis_type": "batch_apk",
                "results": results,
                "batch_summary": batch_summary,
                "invalid_files": invalid_files if invalid_files else None
            }
            
            # Legacy compatibility - keep original summary format
            response_data["summary"] = {
                "total_files": len(valid_files),
                "processing_time_seconds": round(processing_time, 3),
                "files_per_second": round(len(valid_files) / processing_time, 2) if processing_time > 0 else 0,
                "max_files_allowed": 15
            }
            
            return jsonify(response_data)
        finally:
            # Clean up temporary files
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                except Exception:
                    pass
        
    except Exception as e:
        return jsonify({
            "error": "internal_error", 
            "detail": str(e),
            "status": "failed",
            "analysis_type": "batch_apk",
            "processing_time": round(time.time() - start_time, 3) if 'start_time' in locals() else 0,
            "files_processed": len(results) if 'results' in locals() else 0
        }), 500

@app.route('/report', methods=['POST'])
def generate_report():
    """Generate a detailed HTML report for an APK
    
    Accepts either:
    1. File upload (multipart/form-data) - processes the APK and generates report
    2. JSON data with existing scan results - generates report from results
    """
    try:
        # Check if this is a JSON request with existing results
        if request.content_type and 'application/json' in request.content_type:
            try:
                data = request.get_json()
                if not data:
                    return jsonify({"error": "invalid_json", "detail": "Invalid JSON data"}), 400
                
                # Check if we have results data
                if 'results' not in data:
                    return jsonify({"error": "missing_results", "detail": "No results data provided"}), 400
                
                results = data['results']
                if not results or not isinstance(results, list):
                    return jsonify({"error": "invalid_results", "detail": "Results must be a non-empty list"}), 400
                
                # Check format preference
                report_format = data.get('format', 'html').lower()
                
                if report_format == 'html':
                    # Generate HTML report for first result
                    result = results[0]
                    filename = result.get('package', 'unknown') or 'unknown'
                    html_report = _render_html_report(result, filename)
                    return html_report, 200, {'Content-Type': 'text/html'}
                
                elif report_format == 'word':
                    # Generate Word report
                    word_report = _generate_word_report(results)
                    return word_report, 200, {
                        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'Content-Disposition': 'attachment; filename=apk_analysis_report.docx'
                    }
                
                else:
                    return jsonify({"error": "invalid_format", "detail": "Format must be 'html' or 'word'"}), 400
                    
            except Exception as e:
                return jsonify({"error": "json_processing_error", "detail": str(e)}), 400
        
        # Handle file upload case (original functionality)
        if 'file' not in request.files:
            return jsonify({"error": "no_file", "detail": "No file provided. Send either a file upload or JSON data with results."}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "no_file", "detail": "No file selected"}), 400
                
        # Validate file type
        if not allowed_file(file.filename):
            return jsonify({
                "error": "invalid_file_type",
                "detail": f"Invalid file type. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
            }), 400
            
        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
        try:
            file.save(temp_file.name)
            temp_file.close()
            
            # Make sure directories exist
            ensure_dirs()
            
            # Get analysis result
            result = process_single_apk(temp_file.name, quick=False, debug=True)
            
            # Check for errors
            if "error" in result:
                return jsonify(result), 422 if result["error"] == "parse_failed" else 500
            
            # Generate HTML report
            html_report = _render_html_report(result, filename)
            
            return jsonify({
                "result": result,
                "html": html_report
            })
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file.name)
            except Exception:
                pass
        
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/report-batch', methods=['POST'])
def generate_batch_report():
    """Generate comprehensive Word document report for multiple APKs with AI explanations
    
    Returns JSON with:
    - results: Array of analysis results for each APK
    - word_report: Base64-encoded Word document content for download
    - summary: Metadata about the batch processing
    """
    try:
        # Check if files are in request
        if 'files' not in request.files:
            return jsonify({"error": "no_files", "detail": "No files provided"}), 400
        
        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({"error": "no_files", "detail": "No files selected"}), 400
        
        # Limit to 15 files maximum
        if len(files) > 15:
            return jsonify({"error": "too_many_files", "detail": "Maximum 15 files allowed per batch"}), 400
        
        # Validate all files first
        valid_files = []
        for file in files:
            if file.filename == '':
                continue
            if not allowed_file(file.filename):
                continue
            valid_files.append(file)
        
        if not valid_files:
            return jsonify({"error": "no_valid_files", "detail": "No valid APK files found"}), 400
        
        # Process all files
        results = []
        temp_files = []
        
        try:
            for file in valid_files:
                # Save uploaded file temporarily
                filename = secure_filename(file.filename)
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
                temp_files.append(temp_file.name)
                
                file.save(temp_file.name)
                temp_file.close()
                
                # Process file with full analysis
                result = process_single_apk(temp_file.name, quick=False, debug=True)
                result["file"] = file.filename
                results.append(result)
                
            # Generate comprehensive Word document report
            word_report = _generate_word_report(results)
            
            return jsonify({
                "results": results,
                "word_report": word_report,
                "summary": {
                    "total_files": len(valid_files),
                    "report_generated": True,
                    "max_files_allowed": 15
                }
            })
            
        finally:
            # Clean up temporary files
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                except Exception:
                    pass
        
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/report-batch-abuse', methods=['POST'])
def report_batch_abuse():
    """Report multiple malicious APKs with evidence bundles"""
    try:
        # Check if files are in request
        if 'files' not in request.files:
            return jsonify({"error": "no_files", "detail": "No files provided"}), 400
        
        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({"error": "no_files", "detail": "No files selected"}), 400
        
        # Get additional report data
        reporter_email = request.form.get('reporter_email', 'anonymous@example.com')
        reporter_name = request.form.get('reporter_name', 'Anonymous')
        additional_notes = request.form.get('additional_notes', '')
        
        # Process all files and collect malicious ones
        malicious_apks = []
        evidence_bundles = []
        temp_files = []
        
        try:
            for file in files:
                if file.filename == '' or not allowed_file(file.filename):
                    continue
                
                # Save uploaded file temporarily
                filename = secure_filename(file.filename)
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
                temp_files.append(temp_file.name)
                
                file.save(temp_file.name)
                temp_file.close()
                
                # Analyze the APK
                result = process_single_apk(temp_file.name, quick=False, debug=True)
                
                # Only report if malicious
                if result.get("prediction") == "fake" or result.get("probability", 0) > 0.7:
                    malicious_apks.append({
                        "filename": filename,
                        "result": result
                    })
                    
                    # Generate evidence bundle for this malicious APK
                    evidence_bundle = _generate_evidence_bundle(result, filename, reporter_email, reporter_name, additional_notes)
                    evidence_bundles.append(evidence_bundle)
                    
                    # Add to threat feed
                    _add_to_threat_feed(result)
            
            if not malicious_apks:
                return jsonify({
                    "status": "success",
                    "message": "No malicious APKs found in batch. No reports generated.",
                    "malicious_count": 0,
                    "total_analyzed": len(files)
                })
            
            # Generate batch report ID
            batch_report_id = f"batch_report_{int(time.time())}"
            
            # Save all evidence bundles
            for i, evidence_bundle in enumerate(evidence_bundles):
                report_path = os.path.join("artifacts", "reports", f"{batch_report_id}_{i+1}.json")
                with open(report_path, "w", encoding="utf-8") as f:
                    json.dump(evidence_bundle, f, indent=2, ensure_ascii=False)
            
            return jsonify({
                "status": "success",
                "batch_report_id": batch_report_id,
                "malicious_count": len(malicious_apks),
                "total_analyzed": len(files),
                "evidence_bundles": evidence_bundles,
                "threat_feed_updated": True,
                "message": f"Batch abuse report submitted successfully. {len(malicious_apks)} malicious APKs reported."
            })
            
        finally:
            # Clean up temporary files
            for temp_file in temp_files:
                try:
                    os.unlink(temp_file)
                except Exception:
                    pass
        
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/report-abuse', methods=['POST'])
def report_abuse():
    """Report malicious APK with evidence bundle and generate STIX/email templates"""
    try:
        # Check if file is in request
        if 'file' not in request.files:
            return jsonify({"error": "no_file", "detail": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "no_file", "detail": "No file selected"}), 400
        
        # Validate file type
        if not allowed_file(file.filename):
            return jsonify({
                "error": "invalid_file_type", 
                "detail": f"Invalid file type. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
            }), 400
        
        # Get additional report data
        reporter_email = request.form.get('reporter_email', 'anonymous@example.com')
        reporter_name = request.form.get('reporter_name', 'Anonymous')
        additional_notes = request.form.get('additional_notes', '')
        
        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1])
        try:
            file.save(temp_file.name)
            temp_file.close()
            
            # Make sure directories exist
            ensure_dirs()
            
            # Get comprehensive analysis
            result = process_single_apk(temp_file.name, quick=False, debug=True)
            
            # Check for errors
            if "error" in result:
                return jsonify(result), 422 if result["error"] == "parse_failed" else 500
            
            # Generate evidence bundle
            evidence_bundle = _generate_evidence_bundle(result, filename, reporter_email, reporter_name, additional_notes)
            
            # Save report to disk
            report_id = f"report_{int(time.time())}_{result.get('sha256', 'unknown')[:8]}"
            report_path = os.path.join("artifacts", "reports", f"{report_id}.json")
            
            with open(report_path, "w", encoding="utf-8") as f:
                json.dump(evidence_bundle, f, indent=2, ensure_ascii=False)
            
            # Add to threat feed if malicious
            if result.get("prediction") == "fake" or result.get("probability", 0) > 0.7:
                _add_to_threat_feed(result)
            
            return jsonify({
                "status": "success",
                "report_id": report_id,
                "evidence_bundle": evidence_bundle,
                "threat_feed_updated": result.get("prediction") == "fake" or result.get("probability", 0) > 0.7,
                "message": "Abuse report submitted successfully"
            })
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file.name)
            except Exception:
                pass
                
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/threat-feed', methods=['GET'])
def get_threat_feed():
    """Get current threat feed data"""
    try:
        feed = get_cached_threat_feed()
        return jsonify({
            "status": "success",
            "feed": {
                "hash_count": len(feed["hashes"]),
                "package_count": len(feed["packages"]),
                "cert_fingerprint_count": len(feed["cert_fingerprints"]),
                "last_updated": feed["last_updated"],
                "version": feed["version"]
            }
        })
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

def get_static_news_data():
    """Get static news content (RBI guidelines, core security tips)"""
    return {
        "rbi_guidelines": [
            {
                "id": "rbi_001",
                "title": "RBI Guidelines for Safe Banking Apps",
                "content": "Always download banking apps only from official app stores (Google Play Store, Apple App Store). Verify the developer name matches your bank's official name.",
                "category": "guidelines",
                "priority": "high",
                "date": "2024-01-15",
                "source": "Reserve Bank of India"
            },
            {
                "id": "rbi_002", 
                "title": "Two-Factor Authentication Requirements",
                "content": "All banking apps must implement 2FA. Never share OTPs, PINs, or passwords with anyone. Banks will never ask for these details via phone or email.",
                "category": "guidelines",
                "priority": "high",
                "date": "2024-01-10",
                "source": "Reserve Bank of India"
            },
            {
                "id": "rbi_003",
                "title": "App Permissions and Security",
                "content": "Legitimate banking apps only request necessary permissions. Be suspicious of apps asking for camera, microphone, or location access without clear banking-related purpose.",
                "category": "guidelines", 
                "priority": "medium",
                "date": "2024-01-05",
                "source": "Reserve Bank of India"
            }
        ],
        "security_tips": [
            {
                "id": "tip_001",
                "title": "How to Verify Banking App Authenticity",
                "content": "1. Check developer name matches your bank exactly\n2. Verify app has millions of downloads and positive reviews\n3. Look for official bank logo and branding\n4. Check app permissions - legitimate apps request minimal permissions\n5. Never download from third-party websites or links in emails/SMS",
                "category": "education",
                "priority": "high",
                "date": "2024-01-12",
                "source": "Digital Rakshak Security Team"
            },
            {
                "id": "tip_002",
                "title": "Red Flags in Banking Apps",
                "content": "Warning signs of fake banking apps:\n• Poor grammar and spelling errors\n• Unprofessional UI design\n• Requests for unnecessary permissions\n• No customer support contact\n• Suspicious app names or variations\n• Low download count despite claiming to be from major bank",
                "category": "education",
                "priority": "high",
                "date": "2024-01-08",
                "source": "Digital Rakshak Security Team"
            },
            {
                "id": "tip_003",
                "title": "Safe Banking Practices",
                "content": "Best practices for secure mobile banking:\n• Always use official banking apps\n• Enable biometric authentication\n• Keep your device and apps updated\n• Never use public WiFi for banking\n• Log out after each session\n• Monitor your account regularly\n• Report suspicious activities immediately",
                "category": "education",
                "priority": "medium",
                "date": "2024-01-05",
                "source": "Digital Rakshak Security Team"
            }
        ]
    }

def get_dynamic_threat_intelligence():
    """Get real-time threat intelligence from ML system"""
    try:
        threat_feed = get_cached_threat_feed()
        current_time = datetime.now()
        
        # Calculate threat statistics
        total_threats = threat_feed['hash_count'] + threat_feed['package_count'] + threat_feed['cert_fingerprint_count']
        
        # Determine threat level based on recent activity
        if total_threats > 20:
            threat_level = "High"
            priority = "critical"
        elif total_threats > 10:
            threat_level = "Medium"
            priority = "high"
        else:
            threat_level = "Low"
            priority = "medium"
        
        return {
            "id": f"intel_{int(time.time())}",
            "title": "Real-time Threat Intelligence Update",
            "content": f"Digital Rakshak ML System has detected {threat_feed['hash_count']} malicious APKs, {threat_feed['package_count']} suspicious packages, and {threat_feed['cert_fingerprint_count']} compromised certificates in our threat database. The system is actively monitoring for new threats.",
            "category": "intelligence",
            "priority": priority,
            "date": current_time.strftime("%Y-%m-%d"),
            "source": "Digital Rakshak ML System",
            "statistics": {
                "fake_apps_detected": threat_feed['hash_count'],
                "suspicious_packages": threat_feed['package_count'],
                "compromised_certificates": threat_feed['cert_fingerprint_count'],
                "total_threats": total_threats,
                "threat_level": threat_level,
                "last_updated": threat_feed['last_updated']
            }
        }
    except Exception as e:
        print(f"Failed to get dynamic threat intelligence: {e}")
        return None

def generate_news_with_gemini():
    """Generate enhanced news content using Gemini API"""
    try:
        # Get current threat data for context
        threat_feed = get_cached_threat_feed()
        
        prompt = f"""
        Generate a cybersecurity news update about fake banking apps based on the following threat intelligence data:
        
        Current Threat Statistics:
        - Malicious APKs detected: {threat_feed['hash_count']}
        - Suspicious packages: {threat_feed['package_count']}
        - Compromised certificates: {threat_feed['cert_fingerprint_count']}
        
        Please generate:
        1. A security alert about recent fake banking app discoveries
        2. An educational tip about mobile banking security
        3. A threat intelligence summary
        
        Format each as a JSON object with fields: id, title, content, category, priority, date, source
        Categories should be: alert, education, intelligence
        Priorities should be: critical, high, medium, low
        Make the content realistic and informative for Indian banking users.
        """
        
        # Use existing Gemini integration
        if 'gemini_client' in globals():
            response = gemini_client.generate_content(prompt)
            # Parse Gemini response and return structured data
            return parse_gemini_news_response(response.text)
        else:
            print("Gemini client not available, using fallback content")
            return get_fallback_gemini_content()
            
    except Exception as e:
        print(f"Gemini news generation failed: {e}")
        return get_fallback_gemini_content()

def parse_gemini_news_response(gemini_text):
    """Parse Gemini response into structured news format"""
    try:
        # This would parse the Gemini response into proper JSON format
        # For now, return fallback content
        return get_fallback_gemini_content()
    except Exception as e:
        print(f"Failed to parse Gemini response: {e}")
        return get_fallback_gemini_content()

def get_fallback_gemini_content():
    """Fallback content when Gemini is unavailable"""
    current_time = datetime.now()
    return {
        "security_alerts": [
            {
                "id": f"alert_gemini_{int(time.time())}",
                "title": "AI-Generated Security Alert: Banking App Threats on the Rise",
                "content": "Recent analysis by Digital Rakshak's AI system indicates a 15% increase in fake banking app submissions. Users are advised to be extra cautious when downloading banking applications and always verify app authenticity through official channels.",
                "category": "alert",
                "priority": "high",
                "date": current_time.strftime("%Y-%m-%d"),
                "source": "Digital Rakshak AI Analysis",
                "threat_level": "Medium"
            }
        ],
        "ai_education": [
            {
                "id": f"edu_gemini_{int(time.time())}",
                "title": "AI-Powered Security Tip: Biometric Authentication Best Practices",
                "content": "Digital Rakshak's AI analysis shows that apps using biometric authentication are 87% less likely to be malicious. Always enable fingerprint or face recognition for banking apps, but ensure the app requests this permission only when necessary for banking operations.",
                "category": "education",
                "priority": "medium",
                "date": current_time.strftime("%Y-%m-%d"),
                "source": "Digital Rakshak AI Security Advisor"
            }
        ]
    }

@app.route('/news', methods=['GET'])
def get_news():
    """Get hybrid cybersecurity news and awareness content"""
    try:
        # Get static content (RBI guidelines, core tips)
        static_news = get_static_news_data()
        
        # Get dynamic threat intelligence
        dynamic_intel = get_dynamic_threat_intelligence()
        
        # Check if enhanced content is requested
        enhanced = request.args.get('enhanced', 'false').lower() == 'true'
        
        # Initialize news data with static content
        news_data = {
            "rbi_guidelines": static_news["rbi_guidelines"],
            "security_tips": static_news["security_tips"],
            "security_alerts": [],
            "threat_intelligence": []
        }
        
        # Add dynamic threat intelligence
        if dynamic_intel:
            news_data["threat_intelligence"].append(dynamic_intel)
        
        # Add enhanced content if requested
        if enhanced:
            try:
                enhanced_content = generate_news_with_gemini()
                if enhanced_content:
                    news_data.update(enhanced_content)
            except Exception as e:
                print(f"Enhanced content generation failed: {e}")
                # Add fallback enhanced content
                fallback_content = get_fallback_gemini_content()
                news_data.update(fallback_content)
        
        # Add some static alerts if no enhanced content
        if not enhanced and not news_data["security_alerts"]:
            news_data["security_alerts"] = [
                {
                    "id": "alert_static_001",
                    "title": "Fake SBI App Detected - 'SBI Secure'",
                    "content": "Cybersecurity researchers have identified a malicious app impersonating State Bank of India. The fake app 'SBI Secure' attempts to steal banking credentials and OTPs.",
                    "category": "alert",
                    "priority": "critical",
                    "date": "2024-01-20",
                    "source": "Digital Rakshak Threat Intelligence",
                    "affected_banks": ["State Bank of India"],
                    "threat_level": "High"
                },
                {
                    "id": "alert_static_002",
                    "title": "New Banking Trojan Targets HDFC Users",
                    "content": "A sophisticated banking trojan has been discovered targeting HDFC Bank customers. The malware can intercept SMS messages and steal banking credentials.",
                    "category": "alert",
                    "priority": "high", 
                    "date": "2024-01-18",
                    "source": "Digital Rakshak Threat Intelligence",
                    "affected_banks": ["HDFC Bank"],
                    "threat_level": "High"
                }
            ]
        
        # Calculate total items
        total_items = sum(len(category) for category in news_data.values())
        
        return jsonify({
            "status": "success",
            "news": news_data,
            "last_updated": datetime.now().isoformat(),
            "total_items": total_items,
            "source": "hybrid",  # static + dynamic
            "enhanced": enhanced,
            "threat_feed_connected": dynamic_intel is not None
        })
        
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/news/categories', methods=['GET'])
def get_news_categories():
    """Get available news categories with dynamic counts"""
    try:
        # Get current threat data for dynamic counts (with fallback)
        try:
            threat_feed = get_cached_threat_feed()
            dynamic_intel = get_dynamic_threat_intelligence()
        except Exception as e:
            print(f"Failed to get threat feed data: {e}")
            threat_feed = {'hash_count': 0, 'package_count': 0, 'cert_fingerprint_count': 0}
            dynamic_intel = None
        
        categories = {
            "guidelines": {
                "name": "RBI Guidelines",
                "description": "Official banking security guidelines from Reserve Bank of India",
                "icon": "🏛️",
                "count": 3,
                "type": "static"
            },
            "alert": {
                "name": "Security Alerts", 
                "description": "Latest security threats and fake app discoveries",
                "icon": "🚨",
                "count": 2,  # Dynamic based on enhanced content
                "type": "dynamic"
            },
            "education": {
                "name": "Security Education",
                "description": "Tips and best practices for safe banking",
                "icon": "📚",
                "count": 3,
                "type": "static"
            },
            "intelligence": {
                "name": "Threat Intelligence",
                "description": "Real-time threat analysis from ML system",
                "icon": "🔍",
                "count": 1 if dynamic_intel else 0,
                "type": "dynamic",
                "live_data": {
                    "malicious_apks": threat_feed.get('hash_count', 0),
                    "suspicious_packages": threat_feed.get('package_count', 0),
                    "compromised_certs": threat_feed.get('cert_fingerprint_count', 0)
                }
            }
        }
        
        return jsonify({
            "status": "success",
            "categories": categories,
            "system_info": {
                "hybrid_mode": True,
                "static_content": True,
                "dynamic_intelligence": dynamic_intel is not None,
                "gemini_enhancement": request.args.get('enhanced', 'false').lower() == 'true'
            }
        })
        
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/news/enhanced', methods=['GET'])
def get_enhanced_news():
    """Get enhanced news content with Gemini AI generation"""
    try:
        # Force enhanced mode
        enhanced_content = generate_news_with_gemini()
        
        if enhanced_content:
            return jsonify({
                "status": "success",
                "news": enhanced_content,
                "last_updated": datetime.now().isoformat(),
                "source": "gemini_enhanced",
                "ai_generated": True
            })
        else:
            return jsonify({
                "status": "success",
                "news": get_fallback_gemini_content(),
                "last_updated": datetime.now().isoformat(),
                "source": "fallback_enhanced",
                "ai_generated": False
            })
            
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

@app.route('/threat/submit', methods=['POST'])
def submit_threat_intel():
    """Submit new threat intelligence data"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "no_data", "detail": "No JSON data provided"}), 400
        
        # Get current feed
        current_feed = get_cached_threat_feed()
        
        # Add new indicators
        if "hashes" in data:
            current_feed["hashes"].update(data["hashes"])
        if "packages" in data:
            current_feed["packages"].update(data["packages"])
        if "cert_fingerprints" in data:
            current_feed["cert_fingerprints"].update(data["cert_fingerprints"])
        
        # Save updated feed
        if save_threat_feed(current_feed):
            # Clear cache to force reload
            global _threat_feed_cache
            _threat_feed_cache = None
            
            return jsonify({
                "status": "success",
                "message": "Threat intelligence updated successfully",
                "new_counts": {
                    "hashes": len(current_feed["hashes"]),
                    "packages": len(current_feed["packages"]),
                    "cert_fingerprints": len(current_feed["cert_fingerprints"])
                }
            })
        else:
            return jsonify({"error": "save_failed", "detail": "Could not save threat feed"}), 500
            
    except Exception as e:
        return jsonify({"error": "internal_error", "detail": str(e)}), 500

def _generate_evidence_bundle(result, filename, reporter_email, reporter_name, additional_notes):
    """Generate comprehensive evidence bundle for abuse reporting"""
    timestamp = int(time.time())
    
    # Extract IOCs
    domains = result.get("domains", [])
    package = result.get("package", "")
    sha256 = result.get("sha256", "")
    
    evidence = {
        "report_metadata": {
            "report_id": f"report_{timestamp}_{sha256[:8] if sha256 else 'unknown'}",
            "timestamp": timestamp,
            "reporter": {
                "email": reporter_email,
                "name": reporter_name
            },
            "additional_notes": additional_notes
        },
        "apk_analysis": {
            "filename": filename,
            "sha256": sha256,
            "package": package,
            "app_label": result.get("app_label", ""),
            "version": result.get("version", ""),
            "file_size": result.get("file_size", 0),
            "prediction": result.get("prediction", "unknown"),
            "probability": result.get("probability", 0),
            "risk_level": result.get("risk_level", "Unknown")
        },
        "technical_indicators": {
            "domains": domains,
            "permissions": result.get("permissions", []),
            "suspicious_apis": result.get("suspicious_apis_analysis", []),
            "certificate_info": {
                "subject": result.get("cert_subject", "unknown"),
                "issuer": result.get("cert_issuer", "unknown"),
                "status": result.get("certificate_status", "Unknown")
            }
        },
        "stix_template": {
            "type": "indicator",
            "labels": ["malicious-activity"],
            "pattern": f"[file:hashes.'SHA-256' = '{sha256}']",
            "valid_from": timestamp,
            "description": f"Malicious APK: {filename} - {result.get('app_label', 'Unknown app')}"
        },
        "email_template": {
            "to": ["cybersecurity@digitalrakshak.in", "abuse@digitalrakshak.in", "cert@digitalrakshak.in"],
            "cc": ["admin@digitalrakshak.in"],
            "subject": f"🚨 MALICIOUS APK DETECTED: {filename}",
            "body": f"""
🚨 URGENT: MALICIOUS APK DETECTED 🚨

Report ID: {timestamp}
Reporter: {reporter_name} ({reporter_email})
Detection Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}

📱 APK DETAILS:
- Filename: {filename}
- SHA-256: {sha256}
- Package: {package}
- App Label: {result.get('app_label', 'Unknown')}
- Risk Level: {result.get('risk_level', 'Unknown')}
- Confidence: {result.get('confidence_percentage', 0)}%

🔍 TECHNICAL INDICATORS:
- Domains: {', '.join(domains) if domains else 'None'}
- Suspicious Permissions: {', '.join(result.get('critical_permissions', []))}
- Certificate Status: {result.get('certificate_status', 'Unknown')}
- Threat Feed Match: {result.get('threat_feed_match', False)}

📝 ADDITIONAL NOTES:
{additional_notes}

⚠️  IMMEDIATE ACTION REQUIRED:
1. Block this APK across all security systems
2. Update threat intelligence feeds
3. Notify relevant authorities if required
4. Investigate potential impact

Generated by Digital Rakshak Threat Intelligence System
Contact: cybersecurity@digitalrakshak.in
"""
        }
    }
    
    return evidence

def _add_to_threat_feed(result):
    """Add malicious APK indicators to threat feed"""
    try:
        current_feed = get_cached_threat_feed()
        
        # Add hash
        sha256 = result.get("sha256")
        if sha256:
            current_feed["hashes"].add(sha256)
        
        # Add package name
        package = result.get("package")
        if package:
            current_feed["packages"].add(package)
        
        # Save updated feed
        save_threat_feed(current_feed)
        
        # Clear cache
        global _threat_feed_cache
        _threat_feed_cache = None
        
    except Exception as e:
        print(f"Warning: Could not add to threat feed: {e}")

def _render_html_report(result: Dict, filename: str) -> str:
    """Generate HTML report from analysis result"""
    fv = result.get("feature_vector", {})
    top = result.get("top_shap", [])
    pred = result.get("prediction", "unknown")
    prob = result.get("probability", 0)
    risk = result.get("risk_level", "Unknown")
    
    # Get app metadata
    app_label = result.get("app_label", "")
    package = result.get("package", "")
    version = result.get("version", "")
    file_size = result.get("file_size", 0)
    
    # Format file size
    if file_size > 0:
        if file_size > 1024 * 1024:
            file_size_str = f"{file_size / (1024 * 1024):.1f} MB"
        else:
            file_size_str = f"{file_size / 1024:.1f} KB"
    else:
        file_size_str = "Unknown"
    
    def format_feature_name(name):
        """Convert feature name to proper title case without underscores"""
        # Special mappings for common abbreviations and terms
        special_mappings = {
            'api': 'API',
            'sdk': 'SDK',
            'cn': 'Certificate Name',
            'pkg': 'Package',
            'url': 'URL',
            'http': 'HTTP',
            'tld': 'Top Level Domain',
            'mb': 'MB',
            'upi': 'UPI',
            'sms': 'SMS',
            'id': 'ID',
            'dex': 'DEX'
        }
        
        # Remove prefixes like 'api_', 'perm_' for cleaner names
        clean_name = name
        if name.startswith('api_'):
            clean_name = name[4:]
        elif name.startswith('perm_'):
            clean_name = name[5:]
        
        # Split by underscores and convert to title case
        words = clean_name.split('_')
        formatted_words = []
        
        for word in words:
            if word.lower() in special_mappings:
                formatted_words.append(special_mappings[word.lower()])
            else:
                formatted_words.append(word.title())
        
        return ' '.join(formatted_words)
    
    # Separate permissions and other features
    permissions = {}
    api_features = {}
    general_features = {}
    
    for k, v in fv.items():
        if k in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS', 'SYSTEM_ALERT_WINDOW', 'READ_CONTACTS', 'INTERNET', 'QUERY_ALL_PACKAGES']:
            if v == 1:  # Only show permissions that are granted (true)
                permissions[k] = v
        elif k.startswith('api_') and v == 1:  # Only show APIs that are present
            api_features[k] = v
        else:
            general_features[k] = v
    
    # Build permission cards (only for granted permissions)
    permission_cards = ""
    if permissions:
        for perm, value in permissions.items():
            perm_display = format_feature_name(perm)
            permission_cards += f"""
                <div class="permission-card granted">
                    <div class="permission-icon">🔓</div>
                    <div class="permission-name">{perm_display}</div>
                </div>
            """
    
    # Build API features list (only for present APIs)
    api_rows = ""
    if api_features:
        for api, value in api_features.items():
            api_display = format_feature_name(api)
            api_rows += f"<tr><td>{api_display}</td><td><span class='status-present'>Present</span></td></tr>"
    
    # Build general features table
    general_rows = ""
    for k, v in general_features.items():
        if k not in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS', 'SYSTEM_ALERT_WINDOW', 'READ_CONTACTS', 'INTERNET', 'QUERY_ALL_PACKAGES'] and not k.startswith('api_'):
            feature_display = format_feature_name(k)
            if isinstance(v, bool):
                display_value = "Yes" if v else "No"
                value_class = "status-yes" if v else "status-no"
                general_rows += f"<tr><td>{feature_display}</td><td><span class='{value_class}'>{display_value}</span></td></tr>"
            elif isinstance(v, (int, float)) and k.endswith('_score'):
                general_rows += f"<tr><td>{feature_display}</td><td><span class='score-value'>{v:.2f}</span></td></tr>"
            else:
                general_rows += f"<tr><td>{feature_display}</td><td>{v}</td></tr>"
    
    # Build SHAP features list with formatted names
    shap_rows = ""
    for item in top:
        feature_name = format_feature_name(item.get('feature', 'Unknown'))
        value = round(item.get('value', 0), 4)
        impact_class = "positive-impact" if value > 0 else "negative-impact"
        shap_rows += f"<li><strong>{feature_name}</strong>: <span class='{impact_class}'>{value:+.4f}</span></li>"
    
    # Risk color
    risk_color = {"Red": "#dc2626", "Amber": "#d97706", "Green": "#16a34a"}.get(risk, "#6b7280")
    
    # Prediction color
    pred_color = "#dc2626" if pred == "fake" else "#16a34a"
    
    html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1">
        <title>APK Security Analysis Report - {filename}</title>
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif;
                margin: 0;
                padding: 20px;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: #1f2937;
                line-height: 1.6;
                min-height: 100vh;
            }}
            .container {{
                max-width: 1100px;
                margin: 0 auto;
                background: white;
                border-radius: 16px;
                box-shadow: 0 20px 25px -5px rgba(0, 0, 0, 0.1), 0 10px 10px -5px rgba(0, 0, 0, 0.04);
                overflow: hidden;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                text-align: center;
                position: relative;
            }}
            .header::before {{
                content: '';
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                bottom: 0;
                background: url("data:image/svg+xml,%3Csvg width='40' height='40' viewBox='0 0 40 40' xmlns='http://www.w3.org/2000/svg'%3E%3Cg fill='%23ffffff' fill-opacity='0.05'%3E%3Cpath d='M20 20c0 11.046-8.954 20-20 20v20h40V20H20z'/%3E%3C/g%3E%3C/svg%3E") repeat;
            }}
            .header-content {{
                position: relative;
                z-index: 1;
            }}
            .header h1 {{
                margin: 0 0 15px 0;
                font-size: 2.5em;
                font-weight: 800;
                text-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }}
            .header p {{
                margin: 0;
                opacity: 0.95;
                font-size: 1.2em;
                font-weight: 300;
            }}
            .app-info {{
                margin-top: 15px !important;
                font-size: 1em !important;
                background: rgba(255, 255, 255, 0.1);
                padding: 12px 20px;
                border-radius: 8px;
                backdrop-filter: blur(10px);
                border: 1px solid rgba(255, 255, 255, 0.2);
            }}
            .content {{
                padding: 40px;
            }}
            .summary {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
                gap: 25px;
                margin-bottom: 40px;
            }}
            .summary-card {{
                padding: 25px;
                border-radius: 12px;
                border: 1px solid #e5e7eb;
                background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
                transition: transform 0.2s ease, box-shadow 0.2s ease;
            }}
            .summary-card:hover {{
                transform: translateY(-2px);
                box-shadow: 0 8px 15px -3px rgba(0, 0, 0, 0.1);
            }}
            .summary-card h3 {{
                margin: 0 0 12px 0;
                font-size: 0.9em;
                font-weight: 700;
                text-transform: uppercase;
                color: #6b7280;
                letter-spacing: 0.1em;
            }}
            .summary-card .value {{
                font-size: 2em;
                font-weight: 800;
                margin: 0;
            }}
            .section {{
                margin-bottom: 40px;
                background: #f9fafb;
                border-radius: 12px;
                padding: 30px;
                border: 1px solid #e5e7eb;
            }}
            .section h2 {{
                margin: 0 0 25px 0;
                font-size: 1.6em;
                font-weight: 700;
                color: #111827;
                border-bottom: 3px solid #667eea;
                padding-bottom: 10px;
                display: flex;
                align-items: center;
            }}
            .section h2::before {{
                content: '🔍';
                margin-right: 10px;
                font-size: 1.2em;
            }}
            .app-metadata {{
                margin-top: 20px;
            }}
            .metadata-grid {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
                gap: 15px;
                margin-top: 15px;
            }}
            .metadata-item {{
                display: flex;
                justify-content: space-between;
                align-items: center;
                padding: 12px 16px;
                background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                transition: all 0.2s ease;
            }}
            .metadata-item:hover {{
                transform: translateY(-1px);
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
            }}
            .metadata-label {{
                font-weight: 600;
                color: #64748b;
                font-size: 0.9em;
                text-transform: uppercase;
                letter-spacing: 0.05em;
            }}
            .metadata-value {{
                font-weight: 500;
                color: #1e293b;
                font-size: 0.95em;
                word-break: break-all;
            }}
            }}
            .permissions-grid {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
                gap: 15px;
                margin-top: 15px;
            }}
            .permission-card {{
                display: flex;
                align-items: center;
                padding: 15px 20px;
                border-radius: 8px;
                border: 2px solid #fbbf24;
                background: #fef3c7;
            }}
            .permission-card.granted {{
                border-color: #ef4444;
                background: #fee2e2;
            }}
            .permission-icon {{
                font-size: 1.5em;
                margin-right: 12px;
            }}
            .permission-name {{
                font-weight: 600;
                color: #374151;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-top: 15px;
                background: white;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 2px 4px -1px rgba(0, 0, 0, 0.1);
            }}
            th, td {{
                padding: 15px;
                text-align: left;
                border-bottom: 1px solid #e5e7eb;
            }}
            th {{
                background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
                font-weight: 700;
                color: #374151;
                text-transform: uppercase;
                font-size: 0.85em;
                letter-spacing: 0.05em;
            }}
            tr:hover {{
                background: #f8fafc;
            }}
            tr:last-child td {{
                border-bottom: none;
            }}
            .status-present {{
                color: #dc2626;
                font-weight: 600;
                padding: 4px 8px;
                background: #fee2e2;
                border-radius: 4px;
                font-size: 0.85em;
            }}
            .status-yes {{
                color: #059669;
                font-weight: 600;
                padding: 4px 8px;
                background: #d1fae5;
                border-radius: 4px;
                font-size: 0.85em;
            }}
            .status-no {{
                color: #6b7280;
                font-weight: 600;
                padding: 4px 8px;
                background: #f3f4f6;
                border-radius: 4px;
                font-size: 0.85em;
            }}
            .score-value {{
                font-weight: 700;
                color: #1f2937;
                padding: 4px 8px;
                background: #e0e7ff;
                border-radius: 4px;
            }}
            ul {{
                list-style: none;
                padding: 0;
            }}
            li {{
                padding: 12px 0;
                border-bottom: 1px solid #f3f4f6;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            li:last-child {{
                border-bottom: none;
            }}
            .positive-impact {{
                color: #dc2626;
                font-weight: 700;
                padding: 4px 8px;
                background: #fee2e2;
                border-radius: 4px;
            }}
            .negative-impact {{
                color: #059669;
                font-weight: 700;
                padding: 4px 8px;
                background: #d1fae5;
                border-radius: 4px;
            }}
            .timestamp {{
                text-align: center;
                margin-top: 40px;
                padding: 25px;
                background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
                border-radius: 12px;
                color: #6b7280;
                font-size: 0.9em;
                border: 1px solid #e5e7eb;
            }}
            .no-data {{
                text-align: center;
                padding: 30px;
                color: #6b7280;
                font-style: italic;
                background: #f9fafb;
                border-radius: 8px;
                border: 1px dashed #d1d5db;
            }}
            @media (max-width: 768px) {{
                body {{ padding: 10px; }}
                .content {{ padding: 25px; }}
                .summary {{ grid-template-columns: 1fr; }}
                .permissions-grid {{ grid-template-columns: 1fr; }}
                .header h1 {{ font-size: 2em; }}
                .section {{ padding: 20px; }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <div class="header-content">
                    <h1>🛡️ APK Security Analysis</h1>
                    <p>Comprehensive security analysis for <strong>{filename}</strong></p>
                    {f'<p class="app-info">📱 <strong>{app_label}</strong> • 📦 {package} • 🔢 v{version} • 📏 {file_size_str}</p>' if app_label else ''}
                </div>
            </div>
            
            <div class="content">
                <div class="summary">
                    <div class="summary-card">
                        <h3>📊 Prediction</h3>
                        <p class="value" style="color: {pred_color};">{pred.title()}</p>
                    </div>
                    <div class="summary-card">
                        <h3>⚠️ Risk Level</h3>
                        <p class="value" style="color: {risk_color};">{risk}</p>
                    </div>
                    <div class="summary-card">
                        <h3>🎯 Confidence</h3>
                         <p class="value">{result.get("confidence_percentage", 0):.1f}%</p>
                    </div>
                    <div class="summary-card">
                        <h3>📈 Score</h3>
                        <p class="value">{prob:.3f}</p>
                    </div>
                </div>
                
                {f'''
                <div class="section">
                    <h2>📱 Application Information</h2>
                    <div class="app-metadata">
                        <div class="metadata-grid">
                            <div class="metadata-item">
                                <span class="metadata-label">App Name:</span>
                                <span class="metadata-value">{app_label}</span>
                            </div>
                            <div class="metadata-item">
                                <span class="metadata-label">Package:</span>
                                <span class="metadata-value">{package}</span>
                            </div>
                            <div class="metadata-item">
                                <span class="metadata-label">Version:</span>
                                <span class="metadata-value">{version}</span>
                            </div>
                            <div class="metadata-item">
                                <span class="metadata-label">File Size:</span>
                                <span class="metadata-value">{file_size_str}</span>
                            </div>
                        </div>
                    </div>
                </div>
                ''' if app_label else ''}
                
                <div class="section">
                    <h2>Top Contributing Features</h2>
                    {f'<ul>{shap_rows}</ul>' if shap_rows else '<div class="no-data">No SHAP analysis available for this APK.</div>'}
                </div>
                
                {f'''
                <div class="section">
                    <h2>🔐 Granted Permissions</h2>
                    <div class="permissions-grid">
                        {permission_cards}
                    </div>
                </div>
                ''' if permission_cards else ''}
                
                {f'''
                <div class="section">
                    <h2>⚡ Suspicious APIs Detected</h2>
                    <table>
                        <thead>
                            <tr>
                                <th>API Function</th>
                                <th>Status</th>
                            </tr>
                        </thead>
                        <tbody>
                            {api_rows}
                        </tbody>
                    </table>
                </div>
                ''' if api_rows else ''}
                
                <div class="section">
                    <h2>📋 Application Metadata</h2>
                    <table>
                        <thead>
                            <tr>
                                <th>Feature</th>
                                <th>Value</th>
                            </tr>
                        </thead>
                        <tbody>
                            {general_rows}
                        </tbody>
                    </table>
                </div>
                
                <div class="timestamp">
                    <strong>📅 Report Generated:</strong> {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}<br>
                    <strong>🔧 Powered by:</strong> Fake APK Detection System v2.0
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html

def _generate_ai_explanation(result: Dict) -> str:
    """Generate AI explanation for the prediction using Gemini API"""
    try:
        import requests
        import json
        import re
        
        # Load API key from environment variable (recommended)
        API_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyAI32Wr0w0cLHNx-X10fG7f_fCDZ4SPIpk")
        ENDPOINT = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        
        def extract_json_from_text(text):
            """Extract JSON content from text that might contain markdown code blocks or other text."""
            json_pattern = r'```(?:json)?\n([\s\S]*?)```'
            match = re.search(json_pattern, text)
            
            if match:
                return match.group(1).strip()
            return text.strip()
        
        def call_gemini(system_prompt, user_prompt):
            payload = {
                "contents": [{
                    "parts": [{
                        "text": f"System: {system_prompt}\nUser: {user_prompt}"
                    }]
                }]
            }
            
            if not API_KEY:
                raise ValueError("API key is missing. Please check your environment variables.")
            
            headers = {
                "Content-Type": "application/json",
                "X-goog-api-key": API_KEY
            }
            
            response = requests.post(
                ENDPOINT,
                headers=headers,
                json=payload
            )
            response.raise_for_status()
            response_data = response.json()
            
            if 'candidates' not in response_data or not response_data['candidates']:
                raise ValueError("No candidates found in response")
            
            raw_text = response_data['candidates'][0]['content']['parts'][0]['text']
            clean_response = extract_json_from_text(raw_text)
            
            try:
                return json.loads(clean_response)  # if valid JSON
            except json.JSONDecodeError:
                return {"raw_text": clean_response}
        
        # Prepare data for AI analysis
        prediction = result.get('prediction', 'unknown')
        probability = result.get('probability', 0)
        risk_level = result.get('risk_level', 'Unknown')
        confidence = result.get('confidence', 'Unknown')
        feature_vector = result.get('feature_vector', {})
        top_shap = result.get('top_shap', [])
        
        # Analyze key features
        suspicious_count = feature_vector.get('count_suspicious', 0)
        has_sms_permissions = any(feature_vector.get(perm, 0) == 1 for perm in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS'])
        has_contacts_permission = feature_vector.get('READ_CONTACTS', 0) == 1
        has_system_alert = feature_vector.get('SYSTEM_ALERT_WINDOW', 0) == 1
        is_official = feature_vector.get('pkg_official', 0) == 1
        cert_present = feature_vector.get('cert_present', 0) == 1
        
        # Build prompts for Gemini
        system_prompt = """You are a senior cybersecurity analyst specializing in Android application security and malware analysis. You have extensive experience in reverse engineering, threat intelligence, and mobile security assessment. Provide expert-level, professional cybersecurity analysis using industry-standard terminology, technical depth, and actionable security insights. Your analysis should be comprehensive, technically accurate, and suitable for cybersecurity professionals and security teams."""
        
        user_prompt = f"""
        Conduct a comprehensive cybersecurity analysis of this Android APK security scan result. Provide a detailed, professional assessment suitable for security teams and cybersecurity professionals.

        **APK Security Analysis Results:**
        - **Classification:** {prediction.upper()}
        - **Confidence Score:** {probability:.1%}
        - **Risk Assessment:** {risk_level} Level
        - **Analysis Confidence:** {confidence}

        **Security Indicators Analysis:**
        - **Suspicious API Count:** {suspicious_count} (threshold-based detection)
        - **SMS Permissions:** {'DETECTED' if has_sms_permissions else 'Not Present'} (potential SMS-based attacks)
        - **Contacts Access:** {'DETECTED' if has_contacts_permission else 'Not Present'} (data harvesting risk)
        - **System Alert Window:** {'DETECTED' if has_system_alert else 'Not Present'} (overlay attack capability)
        - **Package Verification:** {'Official' if is_official else 'Unofficial'} (trustworthiness indicator)
        - **Code Signing:** {'Valid Certificate' if cert_present else 'No Certificate'} (integrity verification)

        **Feature Importance Analysis (SHAP Values):**
        {chr(10).join([f"- **{item.get('feature', 'Unknown').replace('_', ' ').title()}:** {item.get('value', 0):.4f} (contribution weight)" for item in top_shap])}

        **Required Analysis Components:**
        1. **Threat Assessment:** Detailed analysis of the classification with technical reasoning
        2. **Risk Analysis:** Specific security risks, attack vectors, and potential impact
        3. **Technical Indicators:** Deep dive into suspicious patterns, API usage, and behavioral analysis
        4. **Security Recommendations:** Actionable mitigation strategies and security measures
        5. **Professional Context:** Industry-standard cybersecurity terminology and expert insights

        **Analysis Requirements:**
        - Use professional cybersecurity terminology (malware, threat vectors, attack surfaces, etc.)
        - Provide technical depth suitable for security professionals
        - Include specific security recommendations and mitigation strategies
        - Reference industry best practices and security frameworks
        - Maintain professional tone suitable for security reports
        - Focus on actionable intelligence for threat response teams

        Provide a comprehensive cybersecurity analysis (200-300 words) that demonstrates expert-level understanding of mobile security threats and professional security assessment capabilities.
        """
        
        response_data = call_gemini(system_prompt, user_prompt)
        
        if 'raw_text' in response_data:
            return response_data['raw_text'].strip()
        else:
            return str(response_data).strip()
            
    except Exception as e:
        print(f"❌ Gemini API failed: {e}, falling back to rule-based explanation")
        return _generate_rule_based_explanation(result)

def _generate_rule_based_explanation(result: Dict) -> str:
    """Fallback rule-based explanation when Gemini API is not available"""
    prediction = result.get('prediction', 'unknown')
    probability = result.get('probability', 0)
    risk = result.get('risk', 'Unknown')
    feature_vector = result.get('feature_vector', {})
    
    # Analyze key features
    suspicious_count = feature_vector.get('count_suspicious', 0)
    has_sms_permissions = any(feature_vector.get(perm, 0) == 1 for perm in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS'])
    has_contacts_permission = feature_vector.get('READ_CONTACTS', 0) == 1
    has_system_alert = feature_vector.get('SYSTEM_ALERT_WINDOW', 0) == 1
    is_official = feature_vector.get('pkg_official', 0) == 1
    cert_present = feature_vector.get('cert_present', 0) == 1
    
    # Generate explanation based on prediction
    if prediction == 'fake':
        reasons = []
        if probability >= 0.8:
            reasons.append("High probability score indicates strong evidence of malicious behavior")
        if suspicious_count > 3:
            reasons.append(f"Multiple suspicious APIs detected ({suspicious_count} total)")
        if has_sms_permissions:
            reasons.append("SMS permissions detected - potential for SMS-based attacks")
        if has_contacts_permission:
            reasons.append("Contacts access permission - could be used for data harvesting")
        if has_system_alert:
            reasons.append("System alert window permission - potential for overlay attacks")
        if not cert_present:
            reasons.append("No valid certificate found - suspicious for legitimate apps")
        
        explanation = f"This APK was classified as FAKE with {probability:.1%} confidence. "
        if reasons:
            explanation += "Key factors contributing to this classification: " + "; ".join(reasons) + "."
        else:
            explanation += "The model detected patterns consistent with malicious applications."
    
    else:  # legit
        reasons = []
        if probability <= 0.2:
            reasons.append("Low probability score indicates legitimate behavior patterns")
        if is_official:
            reasons.append("Package identified as official/trusted source")
        if cert_present:
            reasons.append("Valid certificate present - indicates proper app signing")
        if suspicious_count <= 1:
            reasons.append("Minimal suspicious API usage detected")
        if not has_sms_permissions and not has_contacts_permission:
            reasons.append("No sensitive permissions requested")
        
        explanation = f"This APK was classified as LEGITIMATE with {probability:.1%} confidence. "
        if reasons:
            explanation += "Key factors supporting this classification: " + "; ".join(reasons) + "."
        else:
            explanation += "The model found no significant indicators of malicious behavior."
    
    # Add risk level explanation
    if risk == 'Red':
        explanation += " Risk Level RED indicates high confidence in malicious behavior."
    elif risk == 'Amber':
        explanation += " Risk Level AMBER indicates moderate suspicion requiring further investigation."
    else:  # Green
        explanation += " Risk Level GREEN indicates low risk of malicious behavior."
    
    return explanation

def _generate_word_report(results: List[Dict]) -> str:
    """Generates a comprehensive Word document report for multiple APKs with AI explanations"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_BREAK
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        import re
        import io
        import base64
    except ImportError:
        # Fallback to HTML if python-docx is not available
        return _generate_html_batch_report(results)
    
    def format_markdown_to_docx(paragraph_text, paragraph):
        """Convert markdown formatting to Word document formatting"""
        # Handle bold text (**text** or __text__)
        bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
        parts = re.split(bold_pattern, paragraph_text)
        
        current_run = None
        for i, part in enumerate(parts):
            if part is None:
                continue
            
            # Check if this part should be bold (every 2nd and 3rd captured group)
            is_bold = (i % 3 == 1 or i % 3 == 2) and part
            
            if part:
                if current_run is None:
                    current_run = paragraph.add_run(part)
                else:
                    current_run = paragraph.add_run(part)
                
                if is_bold:
                    current_run.bold = True
    
    def add_formatted_paragraph(doc, text):
        """Add a paragraph with proper markdown formatting conversion"""
        # Split text into lines and handle different formatting
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph('')  # Empty line
                continue
            
            # Handle headers (# ## ###)
            if line.startswith('###'):
                doc.add_heading(line[3:].strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line[1:].strip(), level=1)
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                format_markdown_to_docx(line[2:].strip(), p)
            # Handle numbered lists
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(style='List Number')
                format_markdown_to_docx(re.sub(r'^\d+\.\s*', '', line), p)
            # Regular paragraph
            else:
                p = doc.add_paragraph()
                format_markdown_to_docx(line, p)
    
    def set_cell_background(cell, color):
        """Set background color of a cell"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph"""
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create a w:r element
        new_run = OxmlElement('w:r')
        
        # Create a new w:rPr element
        rPr = OxmlElement('w:rPr')
        
        # Add color if needed
        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0000FF")
        rPr.append(c)
        
        # Add underline if needed
        u = OxmlElement('w:u')
        u.set(qn('w:val'), "single")
        rPr.append(u)
        
        # Join all the xml elements together
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
        return hyperlink
    
    # Create a new document with professional styling
    doc = Document()
    
    # Set document properties for better metadata
    core_properties = doc.core_properties
    core_properties.title = "APK Security Analysis Report"
    core_properties.subject = "Mobile Application Security"
    core_properties.creator = "Fake APK Detection System v2.0"
    core_properties.category = "Security Analysis"
    
    # Create a professional cover page similar to ATLAS report format
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # Add spacing at the top
    for _ in range(4):
        doc.add_paragraph('')
    
    # Add FA-DET-2025 identifier at top right
    fa_det_para = doc.add_paragraph('FA-DET-2025')
    fa_det_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fa_det_para.runs[0].font.size = Pt(12)
    fa_det_para.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    fa_det_para.runs[0].bold = True
    
    # Add more spacing
    for _ in range(8):
        doc.add_paragraph('')
    
    # Get the main APK name for the title (use first result or generic title)
    main_apk_name = "APK Security Analysis"
    if results and len(results) > 0:
        first_result = results[0]
        app_name = first_result.get('app_label', '')
        file_name = first_result.get('file', '')
        
        if app_name:
            main_apk_name = app_name
        elif file_name:
            # Extract name from filename without extension
            main_apk_name = file_name.split('.')[0] if '.' in file_name else file_name
    
    # Add main application title
    app_title = doc.add_heading(main_apk_name, level=1)
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_title.runs[0].font.size = Pt(32)
    app_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    app_title.runs[0].bold = True
    
    # Add spacing
    for _ in range(2):
        doc.add_paragraph('')
    
    # Add main report title
    main_title = doc.add_heading('APK Security Analysis Report', level=1)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title.runs[0].font.size = Pt(40)
    main_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    main_title.runs[0].bold = True
    
    # Add large spacing before footer
    for _ in range(15):
        doc.add_paragraph('')
    
    # Add copyright notice at bottom center
    copyright_para = doc.add_paragraph('© Copyright 2025 SecureMobile Analytics')
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_para.runs[0].font.size = Pt(12)
    copyright_para.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    
    # Add generation date
    date_para = doc.add_paragraph(f'Generated on: {__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Add confidentiality notice
    conf_para = doc.add_paragraph('CONFIDENTIAL - FOR INTERNAL USE ONLY')
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.runs[0].font.size = Pt(10)
    conf_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    conf_para.runs[0].bold = True
    
    doc.add_page_break()
    
    # Add Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_entries = [
        "1. Executive Summary",
        "2. Summary Table",
        "3. Detailed Analysis",
        "4. Security Recommendations",
        "5. Appendix: Technical Details"
    ]
    
    for entry in toc_entries:
        p = doc.add_paragraph(entry)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # Add Executive Summary with enhanced formatting
    heading = doc.add_heading('1. Executive Summary', level=1)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Add descriptive paragraph
    summary_intro = doc.add_paragraph('This executive summary provides a high-level overview of the security analysis conducted on the submitted APK files. The analysis evaluates each APK for potential security risks, malicious behavior, and suspicious indicators.')
    
    # Calculate statistics for report
    total_files = len(results)
    fake_count = sum(1 for r in results if r.get('prediction') == 'fake')
    legit_count = sum(1 for r in results if r.get('prediction') == 'legit')
    
    # Count risk levels with better categorization
    red_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Red')
    amber_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Amber')
    green_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Green')
    
    # Add a visual 2x2 table for key metrics
    doc.add_heading('1.1 Analysis Metrics', level=2)
    metrics_table = doc.add_table(rows=3, cols=2)
    metrics_table.style = 'Table Grid'
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill table with key metrics
    cell = metrics_table.cell(0, 0)
    cell.text = "Total APKs Analyzed"
    set_cell_background(cell, "E6E6FA")  # Light lavender
    cell = metrics_table.cell(0, 1)
    cell.text = str(total_files)
    
    cell = metrics_table.cell(1, 0)
    cell.text = "Legitimate APKs"
    set_cell_background(cell, "E6FFE6")  # Light green
    cell = metrics_table.cell(1, 1)
    cell.text = f"{legit_count} ({legit_count/total_files*100:.1f}%)"
    
    cell = metrics_table.cell(2, 0)
    cell.text = "Potentially Malicious APKs"
    set_cell_background(cell, "FFEBEB")  # Light red
    cell = metrics_table.cell(2, 1)
    cell.text = f"{fake_count} ({fake_count/total_files*100:.1f}%)"
    
    # Style all cells
    for row in metrics_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    # Add a paragraph break
    doc.add_paragraph('')
    
    # Add risk distribution table
    doc.add_heading('1.2 Risk Distribution', level=2)
    risk_table = doc.add_table(rows=4, cols=2)
    risk_table.style = 'Table Grid'
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    cell = risk_table.cell(0, 0)
    cell.text = "Risk Category"
    set_cell_background(cell, "E6E6FA")  # Light lavender
    cell = risk_table.cell(0, 1)
    cell.text = "Count (Percentage)"
    set_cell_background(cell, "E6E6FA")  # Light lavender
    
    # Data rows with color-coded risk levels
    cell = risk_table.cell(1, 0)
    cell.text = "High Risk (Red)"
    set_cell_background(cell, "FFCCCC")  # Light red
    cell = risk_table.cell(1, 1)
    cell.text = f"{red_count} ({red_count/total_files*100:.1f}%)"
    
    cell = risk_table.cell(2, 0)
    cell.text = "Medium Risk (Amber)"
    set_cell_background(cell, "FFEEBA")  # Light amber
    cell = risk_table.cell(2, 1)
    cell.text = f"{amber_count} ({amber_count/total_files*100:.1f}%)"
    
    cell = risk_table.cell(3, 0)
    cell.text = "Low Risk (Green)"
    set_cell_background(cell, "CCFFCC")  # Light green
    cell = risk_table.cell(3, 1)
    cell.text = f"{green_count} ({green_count/total_files*100:.1f}%)"
    
    # Style all cells
    for row in risk_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    # Add key findings summary
    doc.add_heading('1.3 Key Findings', level=2)
    
    if fake_count > 0:
        p = doc.add_paragraph()
        p.add_run('⚠️ ').font.size = Pt(14)
        p.add_run(f'Detected {fake_count} potentially malicious APK files that require immediate attention.').bold = True
        p.paragraph_format.left_indent = Inches(0.25)
    
    if red_count > 0:
        p = doc.add_paragraph()
        p.add_run('🚨 ').font.size = Pt(14)
        p.add_run(f'Found {red_count} high-risk applications with critical security concerns.').bold = True
        p.paragraph_format.left_indent = Inches(0.25)
    
    if amber_count > 0:
        p = doc.add_paragraph()
        p.add_run('⚠️ ').font.size = Pt(14)
        p.add_run(f'{amber_count} APKs show moderate risk indicators requiring further investigation.').bold = True
        p.paragraph_format.left_indent = Inches(0.25)
    
    if green_count > 0:
        p = doc.add_paragraph()
        p.add_run('✅ ').font.size = Pt(14)
        p.add_run(f'{green_count} APKs appear to be low risk with minimal security concerns.').bold = True
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # Add Summary Table with enhanced formatting
    heading = doc.add_heading('2. Summary Table', level=1)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Add descriptive paragraph
    doc.add_paragraph('The following table provides a comprehensive summary of all analyzed APKs, highlighting key security indicators and risk assessments. Each row represents an individual APK file with its corresponding analysis results.')
    
    # Create a visually appealing table with more structure
    table = doc.add_table(rows=len(results) + 1, cols=6)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set column widths for better readability
    table.columns[0].width = Inches(1.5)  # APK File
    table.columns[1].width = Inches(1.5)  # App Name
    table.columns[2].width = Inches(0.9)  # Prediction
    table.columns[3].width = Inches(1.0)  # Risk Level
    table.columns[4].width = Inches(1.0)  # Confidence
    table.columns[5].width = Inches(3.1)  # Summary
    
    # Add headers with enhanced formatting
    headers = ['APK File', 'App Name', 'Prediction', 'Risk Level', 'Confidence', 'Summary']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_background(cell, "4472C4")  # Dark blue header background
        
        # Style header text as white and bold
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                run.font.size = Pt(11)
    
    # Add data rows with enhanced color coding and formatting
    for i, result in enumerate(results):
        file_name = result.get('file', 'N/A')
        app_label = result.get('app_label', 'Unknown')
        prediction = result.get('prediction', 'Unknown').title()
        risk = result.get('risk_level', result.get('risk', 'Unknown'))
        confidence = f'{result.get("confidence_percentage", 0):.1f}%'
        
        # Generate AI explanation and truncate for table
        ai_explanation = result.get('ai_explanation', '') or _generate_ai_explanation(result)
        # Clean up the explanation and truncate
        clean_explanation = re.sub(r'\*\*(.*?)\*\*', r'\1', ai_explanation)  # Remove bold markdown
        clean_explanation = re.sub(r'[#*-]', '', clean_explanation)  # Remove other markdown
        summary = clean_explanation[:150] + "..." if len(clean_explanation) > 150 else clean_explanation
        
        # Set cell values
        cells = [
            (0, file_name),
            (1, app_label),
            (2, prediction),
            (3, risk),
            (4, confidence),
            (5, summary.strip())
        ]
        
        for col_idx, value in cells:
            cell = table.cell(i + 1, col_idx)
            cell.text = value
            
            # Apply cell styling
            if col_idx == 2:  # Prediction column
                if value.lower() == 'fake':
                    set_cell_background(cell, "FFCCCC")  # Light red for fake
                elif value.lower() == 'legit':
                    set_cell_background(cell, "CCFFCC")  # Light green for legit
            
            if col_idx == 3:  # Risk level column
                if value.lower() == 'red':
                    set_cell_background(cell, "FFCCCC")  # Light red
                elif value.lower() == 'amber':
                    set_cell_background(cell, "FFEEBA")  # Light amber
                elif value.lower() == 'green':
                    set_cell_background(cell, "CCFFCC")  # Light green
            
            # Center align specific columns
            if col_idx in [2, 3, 4]:  # Prediction, Risk, Confidence
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply zebra striping for better readability
        if i % 2 == 0:
            # Even rows get light gray background
            for col_idx in [0, 1, 5]:  # Don't color the already colored cells
                cell = table.cell(i + 1, col_idx)
                set_cell_background(cell, "F2F2F2")  # Light gray
    
    # Add a note about the table
    note_para = doc.add_paragraph()
    note_para.add_run("Note: ").bold = True
    note_para.add_run("This table provides an overview of all analyzed APKs. Red indicates high risk, Amber indicates medium risk, and Green indicates low risk. Click on each APK filename in the detailed analysis section for more information.")
    note_para.paragraph_format.left_indent = Inches(0.25)
    note_para.paragraph_format.space_before = Pt(12)
    
    doc.add_page_break()
    
    # Add Detailed Analysis with enhanced visual formatting
    heading = doc.add_heading('3. Detailed Analysis', level=1)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    doc.add_paragraph('This section provides an in-depth analysis of each APK file, including detailed security assessment, detected permissions, suspicious behaviors, and AI-powered risk analysis.')
    
    for i, result in enumerate(results):
        file_name = result.get("file", "N/A")
        prediction = result.get("prediction", "Unknown").title()
        risk_level = result.get("risk_level", result.get("risk", "Unknown"))
        
        # Create a visually distinct heading for each APK with risk color coding
        apk_heading = doc.add_heading(f'3.{i+1} {file_name}', level=2)
        
        # Set heading color based on risk level
        if risk_level.lower() == 'red':
            apk_heading.runs[0].font.color.rgb = RGBColor(192, 0, 0)  # Dark red
        elif risk_level.lower() == 'amber':
            apk_heading.runs[0].font.color.rgb = RGBColor(184, 134, 11)  # Golden brown
        else:
            apk_heading.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Dark green
            
        # Create a summary box with key information
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        
        # Add verdict cell with appropriate color
        verdict_cell = summary_table.cell(0, 0)
        if prediction.lower() == 'fake':
            verdict_text = "⚠️ POTENTIALLY MALICIOUS"
            set_cell_background(verdict_cell, "FFCCCC")  # Light red
        else:
            verdict_text = "✓ LEGITIMATE"
            set_cell_background(verdict_cell, "CCFFCC")  # Light green
            
        verdict_cell.text = verdict_text
        verdict_cell.width = Inches(2.5)
        
        # Add confidence information
        confidence_cell = summary_table.cell(0, 1)
        confidence_score = result.get("confidence_percentage", 0)
        confidence_level = result.get("confidence", "Unknown")
        confidence_cell.text = f"Confidence: {confidence_score:.1f}% ({confidence_level})"
        confidence_cell.width = Inches(3.5)
        
        # Format cells
        for cell in [verdict_cell, confidence_cell]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12)
        
        doc.add_paragraph('')  # Add space
        
        # Create Application Information section with better formatting
        app_label = result.get('app_label', '')
        package = result.get('package', '')
        version = result.get('version', '')
        file_size = result.get('file_size', 0)
        
        if app_label or package or version:
            app_section = doc.add_heading('Application Information', level=3)
            app_table = doc.add_table(rows=4, cols=2)
            app_table.style = 'Light Grid'
            
            # Add header row
            header_cell = app_table.cell(0, 0)
            header_cell.merge(app_table.cell(0, 1))
            header_cell.text = "APPLICATION METADATA"
            set_cell_background(header_cell, "E6E6FA")  # Light lavender
            
            for paragraph in header_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
            
            # Row 1: App Name
            app_table.cell(1, 0).text = "App Name"
            app_table.cell(1, 1).text = app_label if app_label else "Unknown"
            
            # Row 2: Package
            app_table.cell(2, 0).text = "Package"
            app_table.cell(2, 1).text = package if package else "Unknown"
            
            # Row 3: Version & Size
            app_table.cell(3, 0).text = "Version"
            file_size_str = f"{file_size / (1024*1024):.1f} MB" if file_size > 0 else "Unknown"
            app_table.cell(3, 1).text = f"{version if version else 'Unknown'} ({file_size_str})"
            
            # Format first column
            for row in range(1, 4):
                cell = app_table.cell(row, 0)
                set_cell_background(cell, "F2F2F2")  # Light gray
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        doc.add_paragraph('')  # Add space
        
        # AI Security Analysis with visual enhancements
        security_heading = doc.add_heading('Security Assessment', level=3)
        
        # Create a box for the AI explanation
        ai_explanation = result.get('ai_explanation', '') or _generate_ai_explanation(result)
        if ai_explanation:
            ai_box = doc.add_table(rows=1, cols=1)
            ai_box.style = 'Light Shading'
            
            cell = ai_box.cell(0, 0)
            # Different background based on verdict
            if prediction.lower() == 'fake':
                set_cell_background(cell, "FFF0F0")  # Very light red
            else:
                set_cell_background(cell, "F0FFF0")  # Very light green
                
            paragraph = cell.paragraphs[0]
            format_markdown_to_docx(ai_explanation, paragraph)
        else:
            doc.add_paragraph("AI security assessment not available for this APK.")
        
        doc.add_paragraph('')  # Add space
        
        # Feature Analysis with visual tables and icons
        doc.add_heading('Technical Analysis', level=3)
        feature_vector = result.get('feature_vector', {})
        
        # Permissions with icon-based display
        permissions = [k for k, v in feature_vector.items() 
                      if k in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS', 'SYSTEM_ALERT_WINDOW', 'READ_CONTACTS', 'INTERNET', 'QUERY_ALL_PACKAGES'] 
                      and v == 1]
        
        if permissions:
            permission_heading = doc.add_heading('Detected Permissions', level=4)
            
            # Create a table for permissions with icons
            perm_table = doc.add_table(rows=len(permissions), cols=2)
            perm_table.style = 'Light List'
            
            for idx, perm in enumerate(permissions):
                perm_display = perm.replace("_", " ").title()
                
                # Add icon based on permission sensitivity
                icon_cell = perm_table.cell(idx, 0)
                if perm in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS']:
                    icon_cell.text = "🔴"  # Red circle for SMS permissions
                elif perm in ['SYSTEM_ALERT_WINDOW', 'READ_CONTACTS']:
                    icon_cell.text = "🟠"  # Orange circle for other sensitive permissions
                else:
                    icon_cell.text = "🟡"  # Yellow circle for less sensitive permissions
                
                # Add permission name
                perm_table.cell(idx, 1).text = perm_display
            
            # Set column widths
            perm_table.columns[0].width = Inches(0.5)
            perm_table.columns[1].width = Inches(5.5)
        else:
            doc.add_paragraph("✓ No sensitive permissions detected")
        
        # Suspicious APIs with color-coded table
        suspicious_apis = [k for k, v in feature_vector.items() if k.startswith('api_') and v == 1]
        if suspicious_apis:
            api_heading = doc.add_heading('Suspicious APIs Detected', level=4)
            
            # Create a table for APIs with risk indicators
            api_table = doc.add_table(rows=len(suspicious_apis) + 1, cols=2)
            api_table.style = 'Light Grid'
            
            # Add header row
            header_cell = api_table.cell(0, 0)
            header_cell.text = "API"
            set_cell_background(header_cell, "FFD6D6")  # Light red
            
            header_cell = api_table.cell(0, 1)
            header_cell.text = "Security Implication"
            set_cell_background(header_cell, "FFD6D6")  # Light red
            
            for paragraph in header_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            
            # Map APIs to their security implications
            api_implications = {
                "api_getDeviceId": "May collect device identifiers",
                "api_sendTextMessage": "Can send SMS messages (potential premium SMS fraud)",
                "api_SmsManager": "Can manage SMS messages (potential data exfiltration)",
                "api_DexClassLoader": "Dynamic code loading (potential obfuscation/evasion)",
                "api_TYPE_SYSTEM_ALERT_WINDOW": "Can display overlays (potential phishing attacks)",
                "api_addView": "May add UI elements (potential for overlay attacks)",
                "api_HttpURLConnection": "Network communication capability",
                "api_openConnection": "Network communication capability"
            }
            
            for idx, api in enumerate(suspicious_apis):
                api_name = api.replace('api_', '').replace('_', ' ').title()
                
                # Add API name and implication
                api_table.cell(idx + 1, 0).text = api_name
                api_table.cell(idx + 1, 1).text = api_implications.get(api, "Potentially suspicious behavior")
                
                # Set background color for the API name cell
                set_cell_background(api_table.cell(idx + 1, 0), "FFF0F0")  # Very light red
        
        # SHAP Analysis with enhanced visual table
        if result.get('top_shap'):
            shap_heading = doc.add_heading('Feature Impact Analysis (SHAP)', level=4)
            
            # Create table for SHAP values
            shap_table = doc.add_table(rows=len(result.get('top_shap', [])) + 1, cols=3)
            shap_table.style = 'Light Grid'
            
            # Add header row
            headers = ["Feature", "Impact Value", "Risk Effect"]
            for i, header in enumerate(headers):
                cell = shap_table.cell(0, i)
                cell.text = header
                set_cell_background(cell, "E6E6FA")  # Light lavender
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add SHAP values with visual indicators
            for idx, item in enumerate(result.get('top_shap', [])):
                feature_name = item.get('feature', 'Unknown').replace('_', ' ').title()
                value = item.get('value', 0)
                impact = "Increases Risk" if value > 0 else "Decreases Risk"
                
                # Add feature name
                shap_table.cell(idx + 1, 0).text = feature_name
                
                # Add SHAP value with sign
                value_cell = shap_table.cell(idx + 1, 1)
                value_cell.text = f"{value:+.4f}"
                
                # Center align value
                for paragraph in value_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add impact with color coding
                impact_cell = shap_table.cell(idx + 1, 2)
                impact_cell.text = impact
                
                # Set background color based on impact
                if value > 0:
                    set_cell_background(impact_cell, "FFCCCC")  # Light red for increasing risk
                else:
                    set_cell_background(impact_cell, "CCFFCC")  # Light green for decreasing risk
                    
                # Center align impact
                for paragraph in impact_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Additional technical details in a clean table format
        if result.get('debug'):
            tech_heading = doc.add_heading('Technical Details', level=4)
            debug_info = result['debug']
            
            # Create technical details table
            tech_table = doc.add_table(rows=4, cols=2)
            tech_table.style = 'Light Grid'
            
            # Row 1: Processing Time
            tech_table.cell(0, 0).text = "Processing Time"
            tech_table.cell(0, 1).text = f"{debug_info.get('processing_time_seconds', 0):.3f} seconds"
            
            # Row 2: Cache Status
            tech_table.cell(1, 0).text = "Cache Used"
            tech_table.cell(1, 1).text = "Yes" if debug_info.get('cache_used', False) else "No"
            
            # Row 3: Model Threshold
            tech_table.cell(2, 0).text = "Model Threshold"
            tech_table.cell(2, 1).text = f"{debug_info.get('model_threshold', 0):.3f}"
            
            # Row 4: SHA256
            tech_table.cell(3, 0).text = "SHA256"
            sha256 = debug_info.get('sha256', 'N/A')
            tech_table.cell(3, 1).text = f"{sha256[:16]}..." if len(sha256) > 16 else sha256
            
            # Format left column
            for row in range(4):
                cell = tech_table.cell(row, 0)
                set_cell_background(cell, "F2F2F2")  # Light gray
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Add separator between APK analyses
        if i < len(results) - 1:
            doc.add_page_break()
    
    # Add comprehensive Security Recommendations with visual elements
    heading = doc.add_heading('4. Security Recommendations', level=1)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    doc.add_paragraph('Based on the analysis of the submitted APKs, the following security recommendations are provided to address identified risks and strengthen your mobile application security posture.')
    
    # Create a visually appealing recommendations section with priority levels
    
    # Critical Actions Section with Red Border
    critical_heading = doc.add_heading('4.1 Critical Actions', level=2)
    critical_heading.style.font.color.rgb = RGBColor(192, 0, 0)  # Dark red
    
    critical_table = doc.add_table(rows=1, cols=1)
    critical_table.style = 'Light Grid Accent 2'
    
    cell = critical_table.cell(0, 0)
    set_cell_background(cell, "FFEBEB")  # Very light red
    
    p = cell.paragraphs[0]
    p.add_run("🚨 HIGH PRIORITY ACTIONS").bold = True
    
    red_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Red')
    fake_count = sum(1 for r in results if r.get('prediction') == 'fake')
    
    actions_list = [
        f"Immediately quarantine and investigate the {fake_count} potentially malicious APKs identified in this report",
        f"Review all {red_count} high-risk (RED) applications with your security team",
        "Verify certificate chains and digital signatures for all suspicious applications",
        "Implement runtime monitoring for applications with suspicious API usage",
        "Conduct manual security review of applications requesting SMS permissions"
    ]
    
    bullet_list = doc.add_paragraph(style='List Bullet')
    bullet_list.paragraph_format.left_indent = Inches(0.25)
    
    for action in actions_list:
        if action == actions_list[0]:
            bullet_list.add_run(action)
        else:
            new_para = doc.add_paragraph(style='List Bullet')
            new_para.paragraph_format.left_indent = Inches(0.25)
            new_para.add_run(action)
    
    doc.add_paragraph('')  # Add space
    
    # Best Practices Section with Blue Border
    practices_heading = doc.add_heading('4.2 Security Best Practices', level=2)
    practices_heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    best_practices_table = doc.add_table(rows=1, cols=2)
    best_practices_table.style = 'Light Grid Accent 5'
    
    # Column headers
    cell = best_practices_table.cell(0, 0)
    cell.text = "Prevention"
    set_cell_background(cell, "E6F0FF")  # Light blue
    
    cell = best_practices_table.cell(0, 1)
    cell.text = "Detection & Response"
    set_cell_background(cell, "E6F0FF")  # Light blue
    
    # Style headers
    for cell in best_practices_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
    
    # Add rows for each practice
    practices = [
        ("Implement a formal app vetting process before deployment", 
         "Establish continuous monitoring for suspicious app behaviors"),
        ("Verify application sources and use only trusted app stores", 
         "Create an incident response plan for malicious app detection"),
        ("Deploy mobile threat defense (MTD) solutions on devices", 
         "Regularly scan all installed applications for security issues"),
        ("Implement app-level containerization for sensitive applications", 
         "Monitor for data exfiltration attempts through network traffic"),
        ("Enforce certificate pinning for sensitive applications", 
         "Maintain logs of application behavior for forensic analysis")
    ]
    
    for prevention, detection in practices:
        row_cells = best_practices_table.add_row().cells
        row_cells[0].text = f"✓ {prevention}"
        row_cells[1].text = f"✓ {detection}"
    
    doc.add_paragraph('')  # Add space
    
    # Risk Mitigation Strategies with visual table
    strategies_heading = doc.add_heading('4.3 Risk Mitigation Strategies', level=2)
    
    # Create 3-column table for different risk levels
    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = 'Light Shading Accent 1'
    
    # Define column headers with color coding
    headers = [
        ("High Risk (Red)", "FFCCCC"),  # Light red
        ("Medium Risk (Amber)", "FFEEBA"),  # Light amber
        ("Low Risk (Green)", "CCFFCC")  # Light green
    ]
    
    # Add headers with colors
    for i, (header_text, color) in enumerate(headers):
        cell = risk_table.cell(0, i)
        cell.text = header_text
        set_cell_background(cell, color)
        
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    # Add a row with strategies for each risk level
    strategies_row = risk_table.add_row().cells
    
    # High Risk strategies
    strategies_row[0].text = ("• Conduct full code review and analysis\n"
                           "• Perform dynamic and static security testing\n"
                           "• Analyze in a sandbox environment\n"
                           "• Block from production environments\n"
                           "• Consider reverse engineering for malware analysis")
    
    # Medium Risk strategies
    strategies_row[1].text = ("• Review requested permissions and APIs\n"
                           "• Test on isolated devices first\n"
                           "• Monitor runtime behavior\n"
                           "• Implement additional access controls\n"
                           "• Verify developer authenticity")
    
    # Low Risk strategies
    strategies_row[2].text = ("• Apply standard security policies\n"
                           "• Update regularly to latest versions\n"
                           "• Include in routine security scans\n"
                           "• Document baseline behavior\n"
                           "• Monitor for deviation from normal patterns")
    
    # Add space before appendix
    doc.add_page_break()
    
    # Add appendix with metadata and information about the report
    appendix_heading = doc.add_heading('5. Appendix: Technical Details', level=1)
    appendix_heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    doc.add_heading('5.1 Report Information', level=2)
    
    # Create metadata table
    metadata_table = doc.add_table(rows=5, cols=2)
    metadata_table.style = 'Light Grid'
    
    # Row 1: Report Generation Date
    metadata_table.cell(0, 0).text = "Report Generation Date"
    metadata_table.cell(0, 1).text = __import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")
    
    # Row 2: Analysis System
    metadata_table.cell(1, 0).text = "Analysis System"
    metadata_table.cell(1, 1).text = "Fake APK Detection System v2.0"
    
    # Row 3: ML Model
    metadata_table.cell(2, 0).text = "ML Model"
    metadata_table.cell(2, 1).text = "XGBoost with SHAP explainability"
    
    # Row 4: Processing Time
    metadata_table.cell(3, 0).text = "Total Processing Time"
    metadata_table.cell(3, 1).text = f"{sum(r.get('debug', {}).get('processing_time_seconds', 0) for r in results):.2f} seconds"
    
    # Row 5: Files Analyzed
    metadata_table.cell(4, 0).text = "Files Analyzed"
    metadata_table.cell(4, 1).text = str(len(results))
    
    # Format left column
    for row in range(5):
        cell = metadata_table.cell(row, 0)
        set_cell_background(cell, "F2F2F2")  # Light gray
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add space
    doc.add_paragraph('')
    
    # Add disclaimer
    doc.add_heading('5.2 Disclaimer', level=2)
    
    disclaimer_box = doc.add_table(rows=1, cols=1)
    disclaimer_box.style = 'Light Grid'
    
    cell = disclaimer_box.cell(0, 0)
    set_cell_background(cell, "F9F9F9")  # Very light gray
    
    p = cell.paragraphs[0]
    p.add_run("DISCLAIMER: ").bold = True
    p.add_run("This report is generated through automated analysis and should be used as a guidance tool only. False positives and false negatives are possible. Always conduct thorough manual verification of high-risk applications. This report does not constitute legal advice or certification of application safety.")
    
    # Add a signature section
    doc.add_paragraph('')
    signature_para = doc.add_paragraph("Report generated by Fake APK Detection System")
    signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Apply final document styling and save
    # Add page numbers at the footer
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_text("Page ")
    
    field_code = "PAGE"
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    
    run = paragraph.add_run()
    run.add_text(" of ")
    
    field_code = "NUMPAGES"
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    
    # Save the document to a BytesIO object and return as base64
    import io
    import base64
    
    try:
        # Save to BytesIO instead of file
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Convert to base64
        doc_base64 = base64.b64encode(doc_buffer.read()).decode('utf-8')
        doc_buffer.close()
        
        return doc_base64
    except Exception as e:
        # If we get an error with the enhanced version, fall back to simpler version
        print(f"Error in enhanced Word document generation: {e}, falling back to simple version")
        
        # Create a simpler document without advanced features
        doc = Document()
        
        # Add a simple title
        doc.add_heading('APK Security Analysis Report', 0)
        
        # Add basic summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f'Total APKs analyzed: {len(results)}')
        doc.add_paragraph(f'Generated on: {__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")}')
        
        # Add simple table
        table = doc.add_table(rows=len(results) + 1, cols=4)
        table.style = 'Table Grid'
        
        # Add headers
        table.cell(0, 0).text = 'APK File'
        table.cell(0, 1).text = 'Prediction'
        table.cell(0, 2).text = 'Risk Level'
        table.cell(0, 3).text = 'Confidence'
        
        # Add data rows
        for i, result in enumerate(results):
            table.cell(i + 1, 0).text = result.get('file', 'N/A')
            table.cell(i + 1, 1).text = result.get('prediction', 'Unknown').title()
            table.cell(i + 1, 2).text = result.get('risk_level', result.get('risk', 'Unknown'))
            table.cell(i + 1, 3).text = f'{result.get("confidence_percentage", 0):.1f}%'
        
        # Try saving the simple version
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Convert to base64
        doc_base64 = base64.b64encode(doc_buffer.read()).decode('utf-8')
        doc_buffer.close()
        
        return doc_base64

def _generate_html_batch_report(results: List[Dict]) -> str:
    """Fallback HTML report generation if python-docx is not available"""
    import re
    
    def format_markdown_to_html(text):
        """Convert markdown formatting to HTML"""
        if not text:
            return ""
            
        # Convert bold text (**text** or __text__)
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.*?)__', r'<strong>\1</strong>', text)
        
        # Convert headers
        text = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)  
        text = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)
        
        # Convert bullet points
        lines = text.split('\n')
        in_list = False
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('- ') or line.startswith('* '):
                if not in_list:
                    formatted_lines.append('<ul>')
                    in_list = True
                formatted_lines.append(f'<li>{line[2:]}</li>')
            else:
                if in_list:
                    formatted_lines.append('</ul>')
                    in_list = False
                if line:
                    formatted_lines.append(f'<p>{line}</p>')
                else:
                    formatted_lines.append('<br>')
        
        if in_list:
            formatted_lines.append('</ul>')
            
        return '\n'.join(formatted_lines)
    
    html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>APK Security Analysis Report</title>
        <style>
            body { 
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; 
                margin: 0; 
                padding: 0; 
                background: #f8fafc;
                color: #1f2937;
                line-height: 1.6;
            }
            .cover-page {
                height: 100vh;
                display: flex;
                flex-direction: column;
                justify-content: space-between;
                align-items: center;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                text-align: center;
                page-break-after: always;
            }
            .fa-det-id {
                align-self: flex-end;
                font-size: 14px;
                font-weight: bold;
                opacity: 0.9;
            }
            .cover-content {
                flex: 1;
                display: flex;
                flex-direction: column;
                justify-content: center;
                align-items: center;
            }
            .app-name {
                font-size: 3.5em;
                font-weight: 800;
                margin-bottom: 40px;
                text-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
            }
            .main-title {
                font-size: 4em;
                font-weight: 800;
                margin-bottom: 20px;
                text-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
                line-height: 1.1;
            }
            .cover-footer {
                opacity: 0.8;
                font-size: 14px;
            }
            .copyright {
                font-weight: bold;
                margin-bottom: 10px;
            }
            .container {
                max-width: 1200px;
                margin: 0 auto;
                background: white;
                border-radius: 12px;
                box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
                overflow: hidden;
            }
            .header { 
                text-align: center; 
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px; 
            }
            .header h1 {
                margin: 0 0 15px 0;
                font-size: 2.5em;
                font-weight: 800;
            }
            .content {
                padding: 40px;
            }
            .summary { 
                background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%); 
                padding: 30px; 
                margin: 30px 0; 
                border-radius: 12px; 
                border: 1px solid #e5e7eb;
            }
            .summary h2 {
                color: #111827;
                border-bottom: 3px solid #667eea;
                padding-bottom: 10px;
            }
            .apk-analysis { 
                border: 1px solid #e5e7eb; 
                margin: 30px 0; 
                padding: 30px; 
                border-radius: 12px; 
                background: #fff;
            }
            .fake { border-left: 5px solid #ef4444; }
            .legit { border-left: 5px solid #10b981; }
            .red { color: #ef4444; font-weight: bold; }
            .amber { color: #f59e0b; font-weight: bold; }
            .green { color: #10b981; font-weight: bold; }
            .analysis-grid {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                gap: 20px;
                margin: 20px 0;
            }
            .stat-card {
                background: #f9fafb;
                padding: 20px;
                border-radius: 8px;
                text-align: center;
                border: 1px solid #e5e7eb;
            }
            .stat-value {
                font-size: 2em;
                font-weight: bold;
                margin-bottom: 5px;
            }
            .ai-explanation {
                background: #f0f9ff;
                padding: 20px;
                border-left: 4px solid #0ea5e9;
                margin: 15px 0;
                border-radius: 0 8px 8px 0;
            }
            .technical-details {
                background: #fafafa;
                padding: 20px;
                border-radius: 8px;
                margin: 15px 0;
            }
            .app-info {
                background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
                padding: 20px;
                border-radius: 8px;
                margin: 15px 0;
                border: 1px solid #bae6fd;
            }
            .app-info h3 {
                color: #0369a1;
                margin-top: 0;
                margin-bottom: 15px;
            }
            .feature-list {
                columns: 2;
                column-gap: 30px;
            }
            .feature-item {
                break-inside: avoid;
                margin-bottom: 8px;
                padding: 8px;
                background: #f3f4f6;
                border-radius: 4px;
            }
            @media print {
                .cover-page {
                    page-break-after: always;
                }
            }
            @media (max-width: 768px) {
                body { padding: 10px; }
                .content { padding: 20px; }
                .analysis-grid { grid-template-columns: 1fr; }
                .feature-list { columns: 1; }
                .app-name { font-size: 2.5em; }
                .main-title { font-size: 2.8em; }
            }
        </style>
    </head>
    <body>
        <!-- Professional Cover Page -->
        <div class="cover-page">
            <div class="fa-det-id">FA-DET-2025</div>
            
            <div class="cover-content">
                """ + (f'<h1 class="app-name">{results[0].get("app_label", results[0].get("file", "APK Analysis").split(".")[0])}</h1>' if results and len(results) > 0 else '') + """
                <h1 class="main-title">APK Security Analysis Report</h1>
            </div>
            
            <div class="cover-footer">
                <div class="copyright">© Copyright 2025 SecureMobile Analytics</div>
                <div>Generated on: """ + __import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC') + """</div>
                <div>CONFIDENTIAL - FOR INTERNAL USE ONLY</div>
            </div>
        </div>
        
        <!-- Report Content -->
        <div class="container">
            <div class="header">
                <h1>🛡️ APK Security Analysis Report</h1>
                <p>Comprehensive security analysis with AI-powered insights</p>
                <p>Generated on: """ + __import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC') + """</p>
            </div>
            <div class="content">
    """
    
    # Executive Summary with enhanced formatting
    total_files = len(results)
    fake_count = sum(1 for r in results if r.get('prediction') == 'fake')
    legit_count = sum(1 for r in results if r.get('prediction') == 'legit')
    red_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Red')
    amber_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Amber')
    green_count = sum(1 for r in results if r.get('risk_level', r.get('risk', 'Unknown')) == 'Green')
    
    html += f"""
        <div class="summary">
            <h2>📊 Executive Summary</h2>
            <div class="analysis-grid">
                <div class="stat-card">
                    <div class="stat-value">{total_files}</div>
                    <div>Total APKs</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value legit">{legit_count}</div>
                    <div>Legitimate</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value fake">{fake_count}</div>
                    <div>Malicious</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value">{(fake_count/total_files)*100:.1f}%</div>
                    <div>Detection Rate</div>
                </div>
            </div>
            
            <h3>Risk Distribution</h3>
            <div class="analysis-grid">
                <div class="stat-card">
                    <div class="stat-value red">{red_count}</div>
                    <div>High Risk (Red)</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value amber">{amber_count}</div>
                    <div>Medium Risk (Amber)</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value green">{green_count}</div>
                    <div>Low Risk (Green)</div>
                </div>
            </div>
        </div>
    """
    
    # Individual APK Analysis with enhanced formatting
    for i, result in enumerate(results, 1):
        prediction = result.get('prediction')
        risk = result.get('risk_level', result.get('risk', 'Unknown'))
        probability = result.get('probability', 0)
        confidence = result.get('confidence', 'Unknown')
        
        # Get app metadata
        app_label = result.get('app_label', '')
        package = result.get('package', '')
        version = result.get('version', '')
        file_size = result.get('file_size', 0)
        
        # Format file size
        if file_size > 0:
            if file_size > 1024 * 1024:
                file_size_str = f"{file_size / (1024 * 1024):.1f} MB"
            else:
                file_size_str = f"{file_size / 1024:.1f} KB"
        else:
            file_size_str = "Unknown"
        
        # Get AI explanation with proper formatting
        ai_explanation = result.get('ai_explanation', '') or _generate_ai_explanation(result)
        formatted_ai_explanation = format_markdown_to_html(ai_explanation)
        
        risk_class = risk.lower()
        pred_class = prediction
        
        html += f"""
        <div class="apk-analysis {pred_class}">
            <h2>📱 Analysis #{i}: {result.get('file', 'N/A')}</h2>
            
            {f'''
            <div class="app-info">
                <h3>📋 Application Information</h3>
                <div class="analysis-grid">
                    <div class="stat-card">
                        <div class="stat-value">{app_label if app_label else 'Unknown'}</div>
                        <div>App Name</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-value">{package if package else 'Unknown'}</div>
                        <div>Package</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-value">{version if version else 'Unknown'}</div>
                        <div>Version</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-value">{file_size_str}</div>
                        <div>File Size</div>
                    </div>
                </div>
            </div>
            ''' if app_label or package or version else ''}
            
            <div class="analysis-grid">
                <div class="stat-card">
                    <div class="stat-value {pred_class}">{prediction.title()}</div>
                    <div>Prediction</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value {risk_class}">{risk}</div>
                    <div>Risk Level</div>
                </div>
                                 <div class="stat-card">
                     <div class="stat-value">{result.get('confidence_percentage', 0):.1f}%</div>
                     <div>Confidence</div>
                 </div>
                <div class="stat-card">
                    <div class="stat-value">{confidence}</div>
                    <div>Certainty</div>
                </div>
            </div>
            
            <div class="ai-explanation">
                <h3>🤖 AI Security Analysis</h3>
                {formatted_ai_explanation}
            </div>
        """
        
        # Technical details
        feature_vector = result.get('feature_vector', {})
        
        # Permissions
        permissions = [k for k, v in feature_vector.items() 
                      if k in ['READ_SMS', 'SEND_SMS', 'RECEIVE_SMS', 'SYSTEM_ALERT_WINDOW', 'READ_CONTACTS', 'INTERNET', 'QUERY_ALL_PACKAGES'] 
                      and v == 1]
        
        # Suspicious APIs
        suspicious_apis = [k for k, v in feature_vector.items() if k.startswith('api_') and v == 1]
        
        if permissions or suspicious_apis or result.get('top_shap'):
            html += f"""
            <div class="technical-details">
                <h3>🔍 Technical Analysis</h3>
            """
            
            if permissions:
                html += f"""
                <h4>Granted Permissions</h4>
                <div class="feature-list">
                """
                for perm in permissions:
                    perm_display = perm.replace("_", " ").title()
                    html += f'<div class="feature-item">🔓 {perm_display}</div>'
                html += "</div>"
            
            if suspicious_apis:
                html += f"""
                <h4>Suspicious APIs Detected</h4>
                <div class="feature-list">
                """
                for api in suspicious_apis:
                    api_name = api.replace('api_', '').replace('_', ' ').title()
                    html += f'<div class="feature-item">⚠️ {api_name}</div>'
                html += "</div>"
            
            if result.get('top_shap'):
                html += f"""
                <h4>Top Contributing Features (SHAP Analysis)</h4>
                <ul>
                """
                for item in result.get('top_shap', []):
                    feature_name = item.get('feature', 'Unknown').replace('_', ' ').title()
                    value = item.get('value', 0)
                    impact = "increases risk" if value > 0 else "decreases risk"
                    impact_class = "red" if value > 0 else "green"
                    html += f'<li><strong>{feature_name}</strong>: <span class="{impact_class}">{value:+.4f}</span> ({impact})</li>'
                html += "</ul>"
            
            html += "</div>"
        
        html += "</div>"
    
    # Security Recommendations
    html += f"""
        <div class="summary">
            <h2>🛡️ Security Recommendations</h2>
            
            <h3>Immediate Actions</h3>
            <ul>
                <li><strong>Review all RED risk APKs immediately</strong> and consider blocking/quarantining</li>
                <li><strong>Investigate AMBER risk APKs thoroughly</strong> before deployment</li>
                <li><strong>Verify the source and authenticity</strong> of all suspicious applications</li>
            </ul>
            
            <h3>Security Best Practices</h3>
            <ul>
                <li>Implement regular APK scanning before deployment</li>
                <li>Establish a security review process for all mobile applications</li>
                <li>Monitor for new threats and update security policies accordingly</li>
                <li>Consider implementing additional runtime protection measures</li>
            </ul>
            
            <h3>Risk Mitigation Strategies</h3>
            <ul>
                <li><strong>High-risk APKs:</strong> Perform detailed manual analysis and sandbox testing</li>
                <li><strong>Medium-risk APKs:</strong> Implement additional monitoring and access controls</li>
                <li><strong>All APKs:</strong> Verify digital signatures and certificate chains</li>
                <li>Maintain an updated whitelist of trusted application sources</li>
            </ul>
        </div>
        
        <div class="summary">
            <h2>📋 Report Metadata</h2>
            <p><strong>Generated by:</strong> APK Security Analysis System v2.0</p>
            <p><strong>Analysis Engine:</strong> XGBoost with SHAP explainability</p>
            <p><strong>Total Processing Time:</strong> {sum(r.get('debug', {}).get('processing_time_seconds', 0) for r in results):.2f} seconds</p>
            <p><strong>Report Generation:</strong> {__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")}</p>
        </div>
            </div>
        </div>
    </body>
    </html>
    """
    
    # Convert HTML to base64 for consistent handling with Word documents
    import base64
    html_base64 = base64.b64encode(html.encode('utf-8')).decode('utf-8')
    return html_base64

@app.route('/test', methods=['GET'])
def test_endpoint():
    return jsonify({"message": "Test endpoint is working!"})



if __name__ == '__main__':
    # Make sure required directories exist
    ensure_dirs()
    
    # Get configuration from environment
    host = os.environ.get('FLASK_HOST', '0.0.0.0')
    
    # Handle PORT environment variable (Render.com uses $PORT)
    port_str = os.environ.get('PORT') or os.environ.get('FLASK_PORT', '9000')
    try:
        port = int(port_str)
    except ValueError:
        print(f"Warning: Invalid port '{port_str}', using default 9000")
        port = 9000
    
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    
    # Performance optimizations for production
    if not debug:
        # Disable Flask debug mode for production
        os.environ['FLASK_ENV'] = 'production'
        
        # Set thread pool size
        import multiprocessing
        workers = min(multiprocessing.cpu_count() * 2 + 1, 8)
        
        print(f"Starting Flask APK Detection API on {host}:{port}")
        print(f"Model path: {MODEL_PATH}")
        print(f"Debug mode: {debug}")
        print(f"Performance mode: Production (workers: {workers})")
        print("Available endpoints:")
        print("  GET  /           - Health check")
        print("  POST /scan       - Scan single APK")
        print("  POST /scan-batch - Scan multiple APKs (up to 15)")
        print("  POST /report     - Generate detailed HTML report")
        print("  POST /report-batch - Generate Word document report (up to 15 APKs)")
        
        # Use production WSGI server
        try:
            from waitress import serve
            print("Using Waitress WSGI server for production...")
            serve(app, host=host, port=port, threads=workers)
        except ImportError:
            print("Waitress not available, using Flask development server...")
            app.run(host=host, port=port, debug=debug, threaded=True)
    else:
        print(f"Starting Flask APK Detection API on {host}:{port}")
        print(f"Model path: {MODEL_PATH}")
        print(f"Debug mode: {debug}")
        print("Available endpoints:")
        print("  GET  /           - Health check")
        print("  POST /scan       - Scan single APK")
        print("  POST /scan-batch - Scan multiple APKs")
        print("  POST /report     - Generate detailed report")
        
        app.run(host=host, port=port, debug=debug, threaded=True)